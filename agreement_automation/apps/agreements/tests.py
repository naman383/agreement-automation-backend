"""Tests for agreement generation."""

from datetime import timedelta
from django.test import TestCase
from django.contrib.auth import get_user_model
from django.utils import timezone
from rest_framework.test import APITestCase, APIClient
from rest_framework import status
from django.urls import reverse

from agreement_automation.apps.templates.models import Template, Placeholder
from .models import UserTemplateUsage, Agreement
from agreement_automation.apps.audit.models import AuditLog

User = get_user_model()


class AgreementTemplateListViewTests(APITestCase):
    """Test cases for AgreementTemplateListView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()
        self.list_url = '/api/v1/agreements/templates/'

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        self.viewer = User.objects.create_user(
            email='viewer@example.com',
            password='viewerpass123'
        )
        self.viewer.role = 'viewer'
        self.viewer.save()

        # Create approved templates
        self.template1 = Template.objects.create(
            name='Content Creator Agreement',
            category='Creator',
            file_path='templates/creator.docx',
            checksum_sha256='abc123',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )
        Placeholder.objects.create(
            template=self.template1,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True
        )

        self.template2 = Template.objects.create(
            name='Dubbing Agreement',
            category='Dubbing',
            file_path='templates/dubbing.docx',
            checksum_sha256='def456',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )
        Placeholder.objects.create(
            template=self.template2,
            name='voice_artist_name',
            display_label='Voice Artist Name',
            field_type='text',
            is_required=True
        )

        # Create draft template (should not appear)
        self.draft_template = Template.objects.create(
            name='Draft Template',
            category='Test',
            file_path='templates/draft.docx',
            checksum_sha256='ghi789',
            status='draft',
            version=1,
            uploaded_by=self.admin
        )

    def test_content_manager_can_list_approved_templates(self):
        """Test that content manager can list approved templates."""
        self.client.force_authenticate(user=self.content_manager)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 2)
        # Should not include draft template
        template_names = [t['name'] for t in response.data['templates']]
        self.assertIn('Content Creator Agreement', template_names)
        self.assertIn('Dubbing Agreement', template_names)
        self.assertNotIn('Draft Template', template_names)

    def test_admin_can_list_approved_templates(self):
        """Test that admin can list approved templates."""
        self.client.force_authenticate(user=self.admin)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 2)

    def test_viewer_cannot_list_templates(self):
        """Test that viewer cannot list templates for agreement generation."""
        self.client.force_authenticate(user=self.viewer)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertIn('Viewers do not have permission', response.data['error'])

    def test_search_by_name(self):
        """Test searching templates by name."""
        self.client.force_authenticate(user=self.content_manager)
        response = self.client.get(f'{self.list_url}?search=dubbing')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 1)
        self.assertEqual(response.data['templates'][0]['name'], 'Dubbing Agreement')

    def test_search_by_category(self):
        """Test searching templates by category."""
        self.client.force_authenticate(user=self.content_manager)
        response = self.client.get(f'{self.list_url}?search=creator')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 1)
        self.assertEqual(response.data['templates'][0]['category'], 'Creator')

    def test_recently_used_templates_appear_first(self):
        """Test that recently used templates appear first."""
        self.client.force_authenticate(user=self.content_manager)

        # Create recent usage for template2
        UserTemplateUsage.objects.create(
            user=self.content_manager,
            template=self.template2,
            usage_count=5
        )

        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        # Recently used template should be first
        self.assertEqual(response.data['templates'][0]['name'], 'Dubbing Agreement')
        self.assertTrue(response.data['templates'][0]['is_recent'])

    def test_old_usage_not_marked_as_recent(self):
        """Test that old usage is not marked as recent."""
        self.client.force_authenticate(user=self.content_manager)

        # Create old usage (8 days ago)
        old_date = timezone.now() - timedelta(days=8)
        usage = UserTemplateUsage.objects.create(
            user=self.content_manager,
            template=self.template2,
            usage_count=5
        )
        UserTemplateUsage.objects.filter(id=usage.id).update(last_used_at=old_date)

        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        # Find template2 in response
        template2_data = next(t for t in response.data['templates'] if t['id'] == self.template2.id)
        self.assertFalse(template2_data['is_recent'])

    def test_template_includes_usage_statistics(self):
        """Test that template includes usage statistics."""
        self.client.force_authenticate(user=self.content_manager)

        # Create usage
        UserTemplateUsage.objects.create(
            user=self.content_manager,
            template=self.template1,
            usage_count=3
        )

        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        template1_data = next(t for t in response.data['templates'] if t['id'] == self.template1.id)
        self.assertEqual(template1_data['user_usage_count'], 3)
        self.assertEqual(template1_data['placeholder_count'], 1)


class AgreementStartViewTests(APITestCase):
    """Test cases for AgreementStartView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()
        self.start_url = '/api/v1/agreements/start/'

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        self.viewer = User.objects.create_user(
            email='viewer@example.com',
            password='viewerpass123'
        )
        self.viewer.role = 'viewer'
        self.viewer.save()

        # Create approved template with placeholders
        self.template = Template.objects.create(
            name='Content Creator Agreement',
            category='Creator',
            file_path='templates/creator.docx',
            checksum_sha256='abc123',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )
        self.placeholder1 = Placeholder.objects.create(
            template=self.template,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True,
            position_index=1
        )
        self.placeholder2 = Placeholder.objects.create(
            template=self.template,
            name='amount',
            display_label='Amount',
            field_type='currency',
            is_required=True,
            position_index=2
        )

        # Create draft template
        self.draft_template = Template.objects.create(
            name='Draft Template',
            category='Test',
            file_path='templates/draft.docx',
            checksum_sha256='xyz789',
            status='draft',
            version=1,
            uploaded_by=self.admin
        )

    def test_content_manager_can_start_agreement(self):
        """Test that content manager can start agreement generation."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('Agreement generation started', response.data['message'])
        self.assertIn('agreement_id', response.data)
        self.assertIn('template', response.data)
        self.assertIn('placeholders', response.data)
        self.assertEqual(len(response.data['placeholders']), 2)

    def test_viewer_cannot_start_agreement(self):
        """Test that viewer cannot start agreement generation."""
        self.client.force_authenticate(user=self.viewer)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_template_usage_recorded(self):
        """Test that template usage is recorded."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        # Verify usage was recorded
        usage = UserTemplateUsage.objects.get(user=self.content_manager, template=self.template)
        self.assertEqual(usage.usage_count, 1)

    def test_template_usage_incremented(self):
        """Test that template usage is incremented on multiple uses."""
        self.client.force_authenticate(user=self.content_manager)

        # Create initial usage
        UserTemplateUsage.objects.create(
            user=self.content_manager,
            template=self.template,
            usage_count=2
        )

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        # Verify usage was incremented
        usage = UserTemplateUsage.objects.get(user=self.content_manager, template=self.template)
        self.assertEqual(usage.usage_count, 3)

    def test_draft_agreement_created(self):
        """Test that draft agreement is created."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        # Verify agreement was created
        agreement = Agreement.objects.get(id=response.data['agreement_id'])
        self.assertEqual(agreement.status, 'draft')
        self.assertEqual(agreement.template, self.template)
        self.assertEqual(agreement.template_version, self.template.version)
        self.assertEqual(agreement.generated_by, self.content_manager)

    def test_cannot_start_with_draft_template(self):
        """Test that agreement cannot be started with draft template."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.draft_template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        self.assertIn('not found or not approved', response.data['error'])

    def test_template_id_required(self):
        """Test that template_id is required."""
        self.client.force_authenticate(user=self.content_manager)

        data = {}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('template_id is required', response.data['error'])

    def test_nonexistent_template(self):
        """Test starting agreement with nonexistent template."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': 99999}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)

    def test_template_without_placeholders(self):
        """Test starting agreement with template without placeholders."""
        self.client.force_authenticate(user=self.content_manager)

        # Create template without placeholders
        template_no_placeholders = Template.objects.create(
            name='Empty Template',
            category='Test',
            file_path='templates/empty.docx',
            checksum_sha256='empty123',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )

        data = {'template_id': template_no_placeholders.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('no placeholders', response.data['error'])

    def test_audit_log_created(self):
        """Test that audit log is created."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        # Verify audit log
        audit_logs = AuditLog.objects.filter(
            user=self.content_manager,
            action='agreement_generation_started'
        )
        self.assertTrue(audit_logs.exists())

    def test_placeholders_returned_in_order(self):
        """Test that placeholders are returned in position_index order."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        placeholders = response.data['placeholders']
        self.assertEqual(placeholders[0]['name'], 'creator_name')
        self.assertEqual(placeholders[1]['name'], 'amount')

    def test_grouped_placeholders_returned(self):
        """Test that grouped_placeholders are returned in response."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('grouped_placeholders', response.data)
        self.assertIsInstance(response.data['grouped_placeholders'], dict)

    def test_grouped_placeholders_by_category(self):
        """Test that placeholders are correctly grouped by category."""
        self.client.force_authenticate(user=self.content_manager)

        # Add more placeholders with different categories
        Placeholder.objects.create(
            template=self.template,
            name='project_name',
            display_label='Project Name',
            field_type='text',
            is_required=True,
            position_index=3
        )
        Placeholder.objects.create(
            template=self.template,
            name='start_date',
            display_label='Start Date',
            field_type='date',
            is_required=True,
            position_index=4
        )

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        grouped = response.data['grouped_placeholders']

        # Check that creator_name is in Creator Information
        self.assertIn('Creator Information', grouped)
        creator_names = [p['name'] for p in grouped['Creator Information']]
        self.assertIn('creator_name', creator_names)

        # Check that amount is in Payment Terms
        self.assertIn('Payment Terms', grouped)
        payment_names = [p['name'] for p in grouped['Payment Terms']]
        self.assertIn('amount', payment_names)

        # Check that project_name is in Project Details
        self.assertIn('Project Details', grouped)
        project_names = [p['name'] for p in grouped['Project Details']]
        self.assertIn('project_name', project_names)

        # Check that start_date is in Dates & Execution
        self.assertIn('Dates & Execution', grouped)
        date_names = [p['name'] for p in grouped['Dates & Execution']]
        self.assertIn('start_date', date_names)

    def test_form_metadata_included(self):
        """Test that form metadata is included in response."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('form_metadata', response.data)
        metadata = response.data['form_metadata']

        self.assertIn('total_fields', metadata)
        self.assertIn('required_fields', metadata)
        self.assertIn('sections', metadata)
        self.assertEqual(metadata['total_fields'], 2)
        self.assertEqual(metadata['required_fields'], 2)
        self.assertIsInstance(metadata['sections'], list)

    def test_backwards_compatibility_flat_placeholders(self):
        """Test that flat placeholders list is still returned for backwards compatibility."""
        self.client.force_authenticate(user=self.content_manager)

        data = {'template_id': self.template.id}
        response = self.client.post(self.start_url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('placeholders', response.data)
        self.assertIsInstance(response.data['placeholders'], list)
        self.assertEqual(len(response.data['placeholders']), 2)


class AgreementFieldValidateViewTests(APITestCase):
    """Test cases for AgreementFieldValidateView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        self.viewer = User.objects.create_user(
            email='viewer@example.com',
            password='viewerpass123'
        )
        self.viewer.role = 'viewer'
        self.viewer.save()

        # Create template with various placeholder types
        self.template = Template.objects.create(
            name='Content Creator Agreement',
            category='Creator',
            file_path='templates/creator.docx',
            checksum_sha256='abc123',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )

        # Create placeholders for different field types
        self.pan_placeholder = Placeholder.objects.create(
            template=self.template,
            name='creator_pan',
            display_label='Creator PAN',
            field_type='PAN',
            is_required=True,
            position_index=1
        )

        self.gst_placeholder = Placeholder.objects.create(
            template=self.template,
            name='company_gst',
            display_label='Company GST',
            field_type='GST',
            is_required=True,
            position_index=2
        )

        self.email_placeholder = Placeholder.objects.create(
            template=self.template,
            name='creator_email',
            display_label='Creator Email',
            field_type='email',
            is_required=True,
            position_index=3
        )

        self.date_placeholder = Placeholder.objects.create(
            template=self.template,
            name='start_date',
            display_label='Start Date',
            field_type='date',
            is_required=True,
            position_index=4
        )

        self.currency_placeholder = Placeholder.objects.create(
            template=self.template,
            name='fee_amount',
            display_label='Fee Amount',
            field_type='currency',
            is_required=True,
            position_index=5
        )

        self.text_placeholder = Placeholder.objects.create(
            template=self.template,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True,
            position_index=6,
            validation_rules={'min_length': 2, 'max_length': 100}
        )

        # Create agreement
        self.agreement = Agreement.objects.create(
            template=self.template,
            template_version=self.template.version,
            generated_by=self.content_manager,
            status='draft',
            placeholder_data={}
        )

    def test_valid_pan_number(self):
        """Test validation of valid PAN number."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_pan',
            'value': 'ABCDE1234F'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])
        self.assertIsNone(response.data['error'])

    def test_invalid_pan_number(self):
        """Test validation of invalid PAN number."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_pan',
            'value': 'ABC123'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertFalse(response.data['is_valid'])
        self.assertIn('PAN should be 10 characters', response.data['error'])

    def test_valid_gst_number(self):
        """Test validation of valid GST number."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'company_gst',
            'value': '22ABCDE1234F1Z5'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])

    def test_invalid_gst_number(self):
        """Test validation of invalid GST number."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'company_gst',
            'value': 'INVALID'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertFalse(response.data['is_valid'])
        self.assertIn('GST should be 15 characters', response.data['error'])

    def test_valid_email(self):
        """Test validation of valid email."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_email',
            'value': 'creator@example.com'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])

    def test_invalid_email(self):
        """Test validation of invalid email."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_email',
            'value': 'not-an-email'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertFalse(response.data['is_valid'])
        self.assertIn('valid email address', response.data['error'])

    def test_valid_date_iso_format(self):
        """Test validation of date in ISO format."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'start_date',
            'value': '2024-01-15'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])
        self.assertEqual(response.data['formatted_value'], '2024-01-15')

    def test_valid_date_dd_mm_yyyy_format(self):
        """Test validation of date in DD/MM/YYYY format."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'start_date',
            'value': '15/01/2024'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])
        self.assertEqual(response.data['formatted_value'], '2024-01-15')

    def test_invalid_date(self):
        """Test validation of invalid date."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'start_date',
            'value': 'not-a-date'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertFalse(response.data['is_valid'])
        self.assertIn('valid date', response.data['error'])

    def test_valid_currency(self):
        """Test validation of currency value."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'fee_amount',
            'value': '50000'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])

    def test_valid_currency_with_symbols(self):
        """Test validation of currency with rupee symbol and commas."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'fee_amount',
            'value': '₹50,000'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])
        self.assertEqual(response.data['formatted_value'], '50000')

    def test_invalid_currency(self):
        """Test validation of invalid currency."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'fee_amount',
            'value': 'not-a-number'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertFalse(response.data['is_valid'])
        self.assertIn('valid numeric amount', response.data['error'])

    def test_valid_text_within_length(self):
        """Test validation of text within min/max length."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_name',
            'value': 'John Doe'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.data['is_valid'])

    def test_text_too_short(self):
        """Test validation of text below min length."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_name',
            'value': 'A'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertFalse(response.data['is_valid'])
        self.assertIn('Minimum length is 2', response.data['error'])

    def test_required_field_empty(self):
        """Test validation of required field with empty value."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_name',
            'value': ''
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertFalse(response.data['is_valid'])
        self.assertIn('required', response.data['error'])

    def test_viewer_cannot_validate(self):
        """Test that viewer cannot validate fields."""
        self.client.force_authenticate(user=self.viewer)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_pan',
            'value': 'ABCDE1234F'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_cannot_validate_others_agreement(self):
        """Test that user cannot validate another user's agreement."""
        other_user = User.objects.create_user(
            email='other@example.com',
            password='otherpass123'
        )
        other_user.role = 'content_manager'
        other_user.save()

        self.client.force_authenticate(user=other_user)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'creator_pan',
            'value': 'ABCDE1234F'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertIn('do not have permission', response.data['error'])

    def test_nonexistent_agreement(self):
        """Test validation with nonexistent agreement."""
        self.client.force_authenticate(user=self.content_manager)

        url = '/api/v1/agreements/99999/validate-field/'
        data = {
            'placeholder_name': 'creator_pan',
            'value': 'ABCDE1234F'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)

    def test_nonexistent_placeholder(self):
        """Test validation with nonexistent placeholder."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'placeholder_name': 'nonexistent_field',
            'value': 'some value'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        self.assertIn('not found in template', response.data['error'])

    def test_placeholder_name_required(self):
        """Test that placeholder_name is required."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/validate-field/'
        data = {
            'value': 'some value'
        }
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('placeholder_name is required', response.data['error'])


class AgreementUpdateDataViewTests(APITestCase):
    """Test cases for AgreementUpdateDataView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        self.viewer = User.objects.create_user(
            email='viewer@example.com',
            password='viewerpass123'
        )
        self.viewer.role = 'viewer'
        self.viewer.save()

        # Create template with placeholders
        self.template = Template.objects.create(
            name='Content Creator Agreement',
            category='Creator',
            file_path='templates/creator.docx',
            checksum_sha256='abc123',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )

        self.placeholder1 = Placeholder.objects.create(
            template=self.template,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True,
            position_index=1
        )

        self.placeholder2 = Placeholder.objects.create(
            template=self.template,
            name='fee_amount',
            display_label='Fee Amount',
            field_type='currency',
            is_required=True,
            position_index=2
        )

        self.placeholder3 = Placeholder.objects.create(
            template=self.template,
            name='project_name',
            display_label='Project Name',
            field_type='text',
            is_required=False,
            position_index=3
        )

        # Create agreement
        self.agreement = Agreement.objects.create(
            template=self.template,
            template_version=self.template.version,
            generated_by=self.content_manager,
            status='draft',
            placeholder_data={}
        )

    def test_update_all_required_fields(self):
        """Test updating all required fields."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('Agreement data updated successfully', response.data['message'])
        self.assertEqual(response.data['fields_updated'], ['creator_name', 'fee_amount'])
        self.assertTrue(response.data['required_fields_complete'])

        # Verify data was saved
        self.agreement.refresh_from_db()
        self.assertIn('creator_name', self.agreement.placeholder_data)
        self.assertEqual(self.agreement.placeholder_data['creator_name']['value'], 'John Doe')
        self.assertEqual(self.agreement.placeholder_data['creator_name']['source'], 'manual')

    def test_update_with_optional_fields(self):
        """Test updating required and optional fields."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000',
                'project_name': 'My Project'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data['total_fields'], 3)

    def test_missing_required_fields(self):
        """Test updating with missing required fields."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe'
                # Missing fee_amount (required)
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('Missing required fields', response.data['error'])
        self.assertIn('Fee Amount', response.data['missing_fields'])

    def test_empty_required_field(self):
        """Test updating with empty required field value."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': '',
                'fee_amount': '50000'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('Missing required fields', response.data['error'])

    def test_invalid_placeholder_name(self):
        """Test updating with placeholder not in template."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'nonexistent_field': 'value'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('not found in template', response.data['error'])

    def test_placeholder_data_required(self):
        """Test that placeholder_data is required."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {}
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('placeholder_data is required', response.data['error'])

    def test_placeholder_data_must_be_dict(self):
        """Test that placeholder_data must be a dictionary."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': 'not a dict'
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('must be a dictionary', response.data['error'])

    def test_viewer_cannot_update(self):
        """Test that viewer cannot update data."""
        self.client.force_authenticate(user=self.viewer)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_cannot_update_others_agreement(self):
        """Test that user cannot update another user's agreement."""
        other_user = User.objects.create_user(
            email='other@example.com',
            password='otherpass123'
        )
        other_user.role = 'content_manager'
        other_user.save()

        self.client.force_authenticate(user=other_user)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_nonexistent_agreement(self):
        """Test updating nonexistent agreement."""
        self.client.force_authenticate(user=self.content_manager)

        url = '/api/v1/agreements/99999/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)

    def test_data_structure_includes_metadata(self):
        """Test that saved data includes metadata."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify metadata structure
        self.agreement.refresh_from_db()
        creator_data = self.agreement.placeholder_data['creator_name']
        self.assertIn('value', creator_data)
        self.assertIn('source', creator_data)
        self.assertIn('modified_at', creator_data)
        self.assertEqual(creator_data['source'], 'manual')

    def test_update_preserves_existing_data(self):
        """Test that updating specific fields preserves other fields."""
        # First update
        self.client.force_authenticate(user=self.content_manager)
        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000',
                'project_name': 'Project 1'
            }
        }
        response = self.client.patch(url, data, format='json')
        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Second update (only update fee_amount and project_name)
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',  # Keep same
                'fee_amount': '60000',  # Change
                'project_name': 'Project 2'  # Change
            }
        }
        response = self.client.patch(url, data, format='json')
        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify all data is present
        self.agreement.refresh_from_db()
        self.assertEqual(self.agreement.placeholder_data['creator_name']['value'], 'John Doe')
        self.assertEqual(self.agreement.placeholder_data['fee_amount']['value'], '60000')
        self.assertEqual(self.agreement.placeholder_data['project_name']['value'], 'Project 2')

    def test_audit_log_created(self):
        """Test that audit log is created."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/update-data/'
        data = {
            'placeholder_data': {
                'creator_name': 'John Doe',
                'fee_amount': '50000'
            }
        }
        response = self.client.patch(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify audit log
        audit_logs = AuditLog.objects.filter(
            user=self.content_manager,
            action='agreement_data_updated'
        )
        self.assertTrue(audit_logs.exists())


class AgreementPreviewGenerateViewTests(APITestCase):
    """Test cases for AgreementPreviewView and AgreementGenerateView."""

    def setUp(self):
        """Set up test client and users."""
        import io
        from docx import Document as DocxDocument
        self.client = APIClient()

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        # Create proper DOCX template with placeholders
        doc = DocxDocument()
        doc.add_paragraph('Content Creator Agreement')
        doc.add_paragraph('Creator Name: {{creator_name}}')
        doc.add_paragraph('Fee Amount: {{fee_amount}}')

        # Save to Django file field
        from django.core.files.base import ContentFile
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Create template
        self.template = Template.objects.create(
            name='Content Creator Agreement',
            category='Creator',
            checksum_sha256='abc123',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )
        self.template.file_path.save('test_template.docx', ContentFile(docx_buffer.getvalue()), save=True)

        # Create placeholders
        Placeholder.objects.create(
            template=self.template,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True,
            position_index=1
        )
        Placeholder.objects.create(
            template=self.template,
            name='fee_amount',
            display_label='Fee Amount',
            field_type='currency',
            is_required=True,
            position_index=2
        )

        # Create agreement with data
        self.agreement = Agreement.objects.create(
            template=self.template,
            template_version=self.template.version,
            generated_by=self.content_manager,
            status='draft',
            placeholder_data={
                'creator_name': {'value': 'John Doe', 'source': 'manual', 'modified_at': '2024-01-15T10:00:00Z'},
                'fee_amount': {'value': '50000', 'source': 'manual', 'modified_at': '2024-01-15T10:00:00Z'}
            }
        )

    def test_preview_success(self):
        """Test successful preview generation."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/preview/'
        response = self.client.post(url, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('Preview generated successfully', response.data['message'])
        self.assertTrue(response.data['preview_available'])

    def test_preview_missing_required_fields(self):
        """Test preview with missing required fields."""
        self.client.force_authenticate(user=self.content_manager)

        # Create agreement without data
        agreement_no_data = Agreement.objects.create(
            template=self.template,
            template_version=self.template.version,
            generated_by=self.content_manager,
            status='draft',
            placeholder_data={}
        )

        url = f'/api/v1/agreements/{agreement_no_data.id}/preview/'
        response = self.client.post(url, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('Missing required fields', response.data['error'])

    def test_generate_success(self):
        """Test successful agreement generation."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/agreements/{self.agreement.id}/generate/'
        response = self.client.post(url, format='json')

        if response.status_code != status.HTTP_201_CREATED:
            print(f"Error: {response.data}")

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('Agreement generated successfully', response.data['message'])
        self.assertTrue(response.data['integrity_verified'])
        self.assertIn('AGR-', response.data['agreement_ref_id'])

        # Verify agreement was updated
        self.agreement.refresh_from_db()
        self.assertEqual(self.agreement.status, 'generated')
        self.assertIsNotNone(self.agreement.agreement_id)
        self.assertIsNotNone(self.agreement.generated_at)
        self.assertTrue(self.agreement.integrity_verified)

    def test_generate_missing_required_fields(self):
        """Test generate with missing required fields."""
        self.client.force_authenticate(user=self.content_manager)

        # Create agreement without data
        agreement_no_data = Agreement.objects.create(
            template=self.template,
            template_version=self.template.version,
            generated_by=self.content_manager,
            status='draft',
            placeholder_data={}
        )

        url = f'/api/v1/agreements/{agreement_no_data.id}/generate/'
        response = self.client.post(url, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('Missing required fields', response.data['error'])

    def test_viewer_cannot_preview(self):
        """Test that viewer cannot generate preview."""
        viewer = User.objects.create_user(email='viewer@example.com', password='viewerpass123')
        viewer.role = 'viewer'
        viewer.save()

        self.client.force_authenticate(user=viewer)

        url = f'/api/v1/agreements/{self.agreement.id}/preview/'
        response = self.client.post(url, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_viewer_cannot_generate(self):
        """Test that viewer cannot generate agreement."""
        viewer = User.objects.create_user(email='viewer@example.com', password='viewerpass123')
        viewer.role = 'viewer'
        viewer.save()

        self.client.force_authenticate(user=viewer)

        url = f'/api/v1/agreements/{self.agreement.id}/generate/'
        response = self.client.post(url, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_cannot_preview_others_agreement(self):
        """Test that user cannot preview another user's agreement."""
        other_user = User.objects.create_user(email='other@example.com', password='otherpass123')
        other_user.role = 'content_manager'
        other_user.save()

        self.client.force_authenticate(user=other_user)

        url = f'/api/v1/agreements/{self.agreement.id}/preview/'
        response = self.client.post(url, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_cannot_generate_others_agreement(self):
        """Test that user cannot generate another user's agreement."""
        other_user = User.objects.create_user(email='other@example.com', password='otherpass123')
        other_user.role = 'content_manager'
        other_user.save()

        self.client.force_authenticate(user=other_user)

        url = f'/api/v1/agreements/{self.agreement.id}/generate/'
        response = self.client.post(url, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)



class AgreementListViewTests(APITestCase):
    """Test cases for AgreementListView (Stories 4.1-4.4)."""
    def setUp(self):
        """Set up test data."""
        import io
        from docx import Document as DocxDocument
        from django.core.files.base import ContentFile

        self.client = APIClient()

        # Create users
        self.admin = User.objects.create_user(email='admin@example.com', password='adminpass123', is_staff=True)
        self.admin.role = 'admin'
        self.admin.save()

        self.legal_reviewer = User.objects.create_user(email='legal@example.com', password='legalpass123')
        self.legal_reviewer.role = 'legal_reviewer'
        self.legal_reviewer.save()

        self.content_manager1 = User.objects.create_user(email='content1@example.com', password='pass123')
        self.content_manager1.role = 'content_manager'
        self.content_manager1.save()

        self.content_manager2 = User.objects.create_user(email='content2@example.com', password='pass123')
        self.content_manager2.role = 'content_manager'
        self.content_manager2.save()

        # Create templates
        doc1 = DocxDocument()
        doc1.add_paragraph('Creator Agreement')
        doc1.add_paragraph('Name: {{creator_name}}')
        doc1.add_paragraph('Project: {{project_name}}')
        docx_buffer1 = io.BytesIO()
        doc1.save(docx_buffer1)
        docx_buffer1.seek(0)

        self.template1 = Template.objects.create(
            name='Content Creator Agreement',
            category='Creator',
            checksum_sha256='abc123',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )
        self.template1.file_path.save('creator.docx', ContentFile(docx_buffer1.getvalue()), save=True)

        doc2 = DocxDocument()
        doc2.add_paragraph('Dubbing Agreement')
        docx_buffer2 = io.BytesIO()
        doc2.save(docx_buffer2)
        docx_buffer2.seek(0)

        self.template2 = Template.objects.create(
            name='Dubbing Agreement',
            category='Dubbing',
            checksum_sha256='def456',
            status='approved',
            version=1,
            uploaded_by=self.admin
        )
        self.template2.file_path.save('dubbing.docx', ContentFile(docx_buffer2.getvalue()), save=True)

        # Create generated agreements
        from django.utils import timezone
        from datetime import timedelta

        # Content manager 1 agreements
        self.agreement1 = Agreement.objects.create(
            agreement_id='AGR-2026-00001',
            template=self.template1,
            template_version=1,
            generated_by=self.content_manager1,
            status='generated',
            placeholder_data={
                'creator_name': {'value': 'John Doe', 'source': 'manual'},
                'project_name': {'value': 'Project Alpha', 'source': 'manual'}
            },
            generated_at=timezone.now() - timedelta(days=1),
            integrity_verified=True
        )

        self.agreement2 = Agreement.objects.create(
            agreement_id='AGR-2026-00002',
            template=self.template2,
            template_version=1,
            generated_by=self.content_manager1,
            status='downloaded',
            placeholder_data={
                'creator_name': {'value': 'Jane Smith', 'source': 'manual'}
            },
            generated_at=timezone.now() - timedelta(days=5),
            downloaded_at=timezone.now() - timedelta(days=4),
            integrity_verified=True
        )

        # Content manager 2 agreements
        self.agreement3 = Agreement.objects.create(
            agreement_id='AGR-2026-00003',
            template=self.template1,
            template_version=1,
            generated_by=self.content_manager2,
            status='generated',
            placeholder_data={
                'creator_name': {'value': 'Bob Wilson', 'source': 'manual'},
                'project_name': {'value': 'Project Beta', 'source': 'manual'}
            },
            generated_at=timezone.now() - timedelta(days=2),
            integrity_verified=True
        )

        # Draft agreement (should not appear in list)
        self.draft_agreement = Agreement.objects.create(
            template=self.template1,
            template_version=1,
            generated_by=self.content_manager1,
            status='draft',
            placeholder_data={}
        )

    def test_content_manager_sees_only_their_agreements(self):
        """Story 4.1: Content manager sees only their own agreements."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 2)

        # Verify only content_manager1's agreements
        agreement_ids = [a['agreement_id'] for a in response.data['agreements']]
        self.assertIn('AGR-2026-00001', agreement_ids)
        self.assertIn('AGR-2026-00002', agreement_ids)
        self.assertNotIn('AGR-2026-00003', agreement_ids)

    def test_agreements_sorted_by_date_descending(self):
        """Story 4.1: Agreements sorted by most recent first."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        # Most recent should be first
        self.assertEqual(response.data['agreements'][0]['agreement_id'], 'AGR-2026-00001')
        self.assertEqual(response.data['agreements'][1]['agreement_id'], 'AGR-2026-00002')

    def test_pagination(self):
        """Story 4.1: Pagination works correctly."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/?page=1&page_size=1')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 1)
        self.assertEqual(response.data['pagination']['page'], 1)
        self.assertEqual(response.data['pagination']['page_size'], 1)
        self.assertEqual(response.data['pagination']['total_count'], 2)
        self.assertEqual(response.data['pagination']['total_pages'], 2)
        self.assertTrue(response.data['pagination']['has_next'])
        self.assertFalse(response.data['pagination']['has_prev'])

    def test_search_by_creator_name(self):
        """Story 4.2: Search agreements by creator name."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/?search=Jane')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 1)
        self.assertEqual(response.data['agreements'][0]['creator_name'], 'Jane Smith')

    def test_search_by_project_name(self):
        """Story 4.2: Search agreements by project name."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/?search=Alpha')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 1)
        self.assertEqual(response.data['agreements'][0]['project_name'], 'Project Alpha')

    def test_search_by_template_name(self):
        """Story 4.2: Search agreements by template name."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/?search=Dubbing')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 1)
        self.assertEqual(response.data['agreements'][0]['template_name'], 'Dubbing Agreement')

    def test_search_no_results(self):
        """Story 4.2: Search with no matches returns empty list."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/?search=NonexistentCreator')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 0)

    def test_filter_by_template_type(self):
        """Story 4.3: Filter agreements by template type."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get(f'/api/v1/agreements/list/?template_id={self.template1.id}')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 1)
        self.assertEqual(response.data['agreements'][0]['template_category'], 'Creator')

    def test_filter_by_date_range(self):
        """Story 4.2: Filter agreements by date range."""
        self.client.force_authenticate(user=self.content_manager1)
        from django.utils import timezone
        from datetime import timedelta

        date_from = (timezone.now() - timedelta(days=3)).isoformat()
        date_to = timezone.now().isoformat()

        response = self.client.get(f'/api/v1/agreements/list/?date_from={date_from}&date_to={date_to}')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 1)
        self.assertEqual(response.data['agreements'][0]['agreement_id'], 'AGR-2026-00001')

    def test_combined_filters(self):
        """Story 4.3: Multiple filters combined with AND logic."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get(f'/api/v1/agreements/list/?template_id={self.template1.id}&search=John')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 1)
        self.assertEqual(response.data['agreements'][0]['creator_name'], 'John Doe')

    def test_legal_reviewer_sees_all_agreements(self):
        """Story 4.4: Legal reviewer sees all organization agreements."""
        self.client.force_authenticate(user=self.legal_reviewer)
        response = self.client.get('/api/v1/agreements/list/')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 3)

        # Should see agreements from both content managers
        agreement_ids = [a['agreement_id'] for a in response.data['agreements']]
        self.assertIn('AGR-2026-00001', agreement_ids)
        self.assertIn('AGR-2026-00002', agreement_ids)
        self.assertIn('AGR-2026-00003', agreement_ids)

    def test_legal_reviewer_sees_generated_by_info(self):
        """Story 4.4: Legal reviewer sees which user generated each agreement."""
        self.client.force_authenticate(user=self.legal_reviewer)
        response = self.client.get('/api/v1/agreements/list/')

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify generated_by info is included
        for agreement_data in response.data['agreements']:
            self.assertIn('generated_by', agreement_data)
            self.assertIsNotNone(agreement_data['generated_by'])
            self.assertIn('email', agreement_data['generated_by'])

    def test_admin_sees_all_agreements(self):
        """Story 4.4: Admin sees all organization agreements."""
        self.client.force_authenticate(user=self.admin)
        response = self.client.get('/api/v1/agreements/list/')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['agreements']), 3)

    def test_draft_agreements_not_included(self):
        """Verify draft agreements are not included in list."""
        self.client.force_authenticate(user=self.content_manager1)
        response = self.client.get('/api/v1/agreements/list/')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        # Should only see 2 generated/downloaded agreements, not the draft
        self.assertEqual(len(response.data['agreements']), 2)
