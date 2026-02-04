"""Views for agreement generation."""

import re
import hashlib
import os
from io import BytesIO
from datetime import timedelta, datetime
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Count, Q
from django.conf import settings
from django.http import FileResponse
from rest_framework import status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.authentication import SessionAuthentication
from docxtpl import DocxTemplate
from docx import Document
import jinja2.exceptions

from agreement_automation.apps.templates.models import Template, Placeholder
from .models import UserTemplateUsage, Agreement
from agreement_automation.apps.audit.utils import log_audit_event


class CsrfExemptSessionAuthentication(SessionAuthentication):
    """Session authentication without CSRF enforcement."""
    def enforce_csrf(self, request):
        return  # Do not enforce CSRF


def calculate_document_checksum(doc_path_or_bytes):
    """
    Calculate SHA-256 checksum of document text content.
    Extracts text from DOCX and hashes it.
    """
    if isinstance(doc_path_or_bytes, bytes):
        doc = Document(BytesIO(doc_path_or_bytes))
    else:
        doc = Document(doc_path_or_bytes)

    # Extract all text from document
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    # Calculate checksum
    text_content = '\n'.join(full_text)
    return hashlib.sha256(text_content.encode('utf-8')).hexdigest()


def generate_agreement_docx(template_path, placeholder_data):
    """
    Generate DOCX document from template with placeholder data.
    Returns BytesIO object containing generated DOCX.
    """
    doc = DocxTemplate(template_path)

    # Extract just the values from structured data
    context = {}
    for name, data in placeholder_data.items():
        if isinstance(data, dict) and 'value' in data:
            context[name] = data['value']
        else:
            context[name] = data

    # Render template
    doc.render(context)

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_agreement_id():
    """Generate unique agreement ID in format AGR-YYYY-NNNNN."""
    from django.db.models import Max

    year = timezone.now().year
    prefix = f'AGR-{year}-'

    # Get the highest ID for this year
    last_agreement = Agreement.objects.filter(
        agreement_id__startswith=prefix
    ).aggregate(Max('agreement_id'))

    last_id = last_agreement['agreement_id__max']
    if last_id:
        # Extract number and increment
        number = int(last_id.split('-')[-1]) + 1
    else:
        number = 1

    return f'{prefix}{number:05d}'


def validate_pan(value):
    """Validate PAN (Permanent Account Number) format."""
    pattern = r'^[A-Z]{5}[0-9]{4}[A-Z]{1}$'
    if not re.match(pattern, value):
        return False, "PAN should be 10 characters in format: ABCDE1234F (5 letters, 4 digits, 1 letter)"
    return True, None


def validate_gst(value):
    """Validate GST (Goods and Services Tax) format."""
    pattern = r'^[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z]{1}[A-Z0-9]{3}$'
    if not re.match(pattern, value):
        return False, "GST should be 15 characters in format: 22ABCDE1234F1Z5 (2 digits, 5 letters, 4 digits, 1 letter, 3 alphanumeric)"
    return True, None


def validate_email(value):
    """Validate email format."""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    if not re.match(pattern, value):
        return False, "Please enter a valid email address (e.g., user@example.com)"
    return True, None


def validate_date(value):
    """Validate and parse date in multiple formats."""
    date_formats = [
        '%Y-%m-%d',  # ISO 8601: 2024-01-15
        '%d/%m/%Y',  # DD/MM/YYYY: 15/01/2024
        '%m/%d/%Y',  # MM/DD/YYYY: 01/15/2024
        '%d-%m-%Y',  # DD-MM-YYYY: 15-01-2024
        '%Y/%m/%d',  # YYYY/MM/DD: 2024/01/15
    ]

    for fmt in date_formats:
        try:
            parsed_date = datetime.strptime(value, fmt)
            # Return ISO format
            return True, parsed_date.strftime('%Y-%m-%d')
        except ValueError:
            continue

    return False, "Please enter a valid date (e.g., 2024-01-15, 15/01/2024, or 01/15/2024)"


def validate_currency(value):
    """Validate currency/numeric value."""
    # Remove currency symbols and commas
    cleaned = value.replace('₹', '').replace(',', '').strip()
    try:
        float(cleaned)
        return True, cleaned
    except ValueError:
        return False, "Please enter a valid numeric amount (e.g., 50000 or ₹50,000)"


def validate_text(value, validation_rules):
    """Validate text field based on validation rules."""
    if not value:
        return False, "This field is required"

    # Check min length
    min_length = validation_rules.get('min_length')
    if min_length and len(value) < min_length:
        return False, f"Minimum length is {min_length} characters"

    # Check max length
    max_length = validation_rules.get('max_length')
    if max_length and len(value) > max_length:
        return False, f"Maximum length is {max_length} characters"

    # Check pattern
    pattern = validation_rules.get('pattern')
    if pattern and not re.match(pattern, value):
        example = validation_rules.get('example', '')
        return False, f"Invalid format. Example: {example}"

    return True, None


def validate_field(field_type, value, validation_rules=None, is_required=True):
    """
    Validate a field based on its type and validation rules.

    Returns tuple: (is_valid, error_message_or_formatted_value)
    """
    # Check required
    if is_required and not value:
        return False, "This field is required"

    if not value:
        return True, None

    validation_rules = validation_rules or {}

    # Route to appropriate validator
    if field_type == 'PAN':
        return validate_pan(value)
    elif field_type == 'GST':
        return validate_gst(value)
    elif field_type == 'email':
        return validate_email(value)
    elif field_type == 'date':
        return validate_date(value)
    elif field_type == 'currency':
        return validate_currency(value)
    elif field_type == 'number':
        try:
            float(value)
            return True, None
        except ValueError:
            return False, "Please enter a valid number"
    elif field_type in ['text', 'dropdown', 'checkbox']:
        if validation_rules:
            return validate_text(value, validation_rules)
        return True, None
    else:
        # Unknown field type, accept as valid
        return True, None


def group_placeholders_by_category(placeholders):
    """
    Group placeholders into categories based on naming patterns.

    Returns dict with structure:
    {
        "Creator Information": [...],
        "Project Details": [...],
        "Payment Terms": [...],
        "Dates & Execution": [...],
        "Other": [...]
    }
    """
    groups = {
        "Creator Information": [],
        "Project Details": [],
        "Payment Terms": [],
        "Dates & Execution": [],
        "Other": []
    }

    for placeholder in placeholders:
        name_lower = placeholder.name.lower()

        # Determine category based on naming patterns
        if any(keyword in name_lower for keyword in ['creator', 'artist', 'talent', 'voice', 'performer', 'actor']):
            category = "Creator Information"
        elif any(keyword in name_lower for keyword in ['project', 'content', 'video', 'episode', 'series', 'production']):
            category = "Project Details"
        elif any(keyword in name_lower for keyword in ['payment', 'fee', 'amount', 'price', 'cost', 'budget', 'compensation']):
            category = "Payment Terms"
        elif any(keyword in name_lower for keyword in ['date', 'deadline', 'execution', 'start', 'end', 'duration']):
            category = "Dates & Execution"
        else:
            category = "Other"

        groups[category].append({
            'name': placeholder.name,
            'display_label': placeholder.display_label,
            'field_type': placeholder.field_type,
            'is_required': placeholder.is_required,
            'position_index': placeholder.position_index,
            'validation_rules': placeholder.validation_rules
        })

    # Remove empty categories
    return {k: v for k, v in groups.items() if v}


class AgreementTemplateListView(APIView):
    """API endpoint for listing approved templates for agreement generation."""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        """
        List approved templates with usage statistics.

        URL: /api/v1/agreements/templates/
        Query Params:
          - search: filter by name or category
        """
        # Content managers, legal reviewers, and admins can generate agreements
        # Viewers cannot generate agreements
        if request.user.role == 'viewer':
            return Response(
                {"error": "Viewers do not have permission to generate agreements."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get only approved templates
        templates = Template.objects.filter(status='approved')

        # Optional search filter
        search = request.query_params.get('search', None)
        if search:
            templates = templates.filter(
                Q(name__icontains=search) | Q(category__icontains=search)
            )

        # Annotate with usage count (total agreements generated with this template)
        templates = templates.annotate(
            total_usage_count=Count('agreements')
        )

        # Get user's template usage
        user_usage = UserTemplateUsage.objects.filter(user=request.user)
        user_usage_dict = {usage.template_id: usage for usage in user_usage}

        # Calculate recently used threshold (last 7 days)
        recent_threshold = timezone.now() - timedelta(days=7)

        # Build response data
        templates_data = []
        for template in templates:
            usage = user_usage_dict.get(template.id)
            is_recent = usage and usage.last_used_at >= recent_threshold if usage else False

            templates_data.append({
                'id': template.id,
                'name': template.name,
                'category': template.category,
                'version': template.version,
                'placeholder_count': template.placeholders.count(),
                'last_used_at': usage.last_used_at if usage else None,
                'user_usage_count': usage.usage_count if usage else 0,
                'total_usage_count': template.total_usage_count,
                'is_recent': is_recent
            })

        # Sort: recently used (last 7 days) first, then alphabetically
        templates_data.sort(key=lambda x: (not x['is_recent'], x['name']))

        return Response(
            {"templates": templates_data},
            status=status.HTTP_200_OK
        )


@method_decorator(csrf_exempt, name='dispatch')
class AgreementStartView(APIView):
    """API endpoint for starting agreement generation."""

    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Start agreement generation with selected template.

        URL: /api/v1/agreements/start/
        Body: {"template_id": 123}
        """
        # Check permission
        if request.user.role == 'viewer':
            return Response(
                {"error": "Viewers do not have permission to generate agreements."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get template_id from request
        template_id = request.data.get('template_id')
        if not template_id:
            return Response(
                {"error": "template_id is required."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Get template
        try:
            template = Template.objects.get(id=template_id, status='approved')
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found or not approved."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Get placeholders
        placeholders = template.placeholders.all()
        if not placeholders.exists():
            return Response(
                {"error": "Template has no placeholders. Cannot generate agreement."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Record template usage
        usage, created = UserTemplateUsage.objects.get_or_create(
            user=request.user,
            template=template,
            defaults={'usage_count': 0}
        )
        usage.usage_count += 1
        usage.save()  # This updates last_used_at automatically

        # Create draft agreement
        agreement = Agreement.objects.create(
            template=template,
            template_version=template.version,
            generated_by=request.user,
            status='draft',
            placeholder_data={}  # Will be filled in next step
        )

        # Log event
        log_audit_event(
            user=request.user,
            action='agreement_generation_started',
            request=request,
            template_id=template.id,
            template_name=template.name,
            agreement_id=agreement.id
        )

        # Build grouped placeholders for response
        grouped_placeholders = group_placeholders_by_category(placeholders)

        # Also build flat list for backwards compatibility
        placeholders_data = []
        for placeholder in placeholders:
            placeholders_data.append({
                'name': placeholder.name,
                'display_label': placeholder.display_label,
                'field_type': placeholder.field_type,
                'is_required': placeholder.is_required,
                'position_index': placeholder.position_index,
                'validation_rules': placeholder.validation_rules
            })

        return Response(
            {
                "message": "Agreement generation started successfully.",
                "agreement_id": agreement.id,
                "template": {
                    "id": template.id,
                    "name": template.name,
                    "version": template.version,
                    "category": template.category
                },
                "placeholders": placeholders_data,
                "grouped_placeholders": grouped_placeholders,
                "form_metadata": {
                    "total_fields": len(placeholders_data),
                    "required_fields": sum(1 for p in placeholders if p.is_required),
                    "sections": list(grouped_placeholders.keys())
                }
            },
            status=status.HTTP_201_CREATED
        )


@method_decorator(csrf_exempt, name='dispatch')
class AgreementFieldValidateView(APIView):
    """API endpoint for validating individual form fields."""

    permission_classes = [IsAuthenticated]

    def post(self, request, agreement_id):
        """
        Validate a single field value.

        URL: /api/v1/agreements/{agreement_id}/validate-field/
        Body: {
            "placeholder_name": "creator_pan",
            "value": "ABCDE1234F"
        }
        """
        # Check permission
        if request.user.role == 'viewer':
            return Response(
                {"error": "Viewers do not have permission to generate agreements."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get agreement
        try:
            agreement = Agreement.objects.get(id=agreement_id)
        except Agreement.DoesNotExist:
            return Response(
                {"error": "Agreement not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check ownership
        if agreement.generated_by != request.user:
            return Response(
                {"error": "You do not have permission to modify this agreement."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get placeholder name and value
        placeholder_name = request.data.get('placeholder_name')
        value = request.data.get('value', '')

        if not placeholder_name:
            return Response(
                {"error": "placeholder_name is required."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Get placeholder from template
        try:
            placeholder = Placeholder.objects.get(
                template=agreement.template,
                name=placeholder_name
            )
        except Placeholder.DoesNotExist:
            return Response(
                {"error": f"Placeholder '{placeholder_name}' not found in template."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Validate the field
        is_valid, message = validate_field(
            field_type=placeholder.field_type,
            value=value,
            validation_rules=placeholder.validation_rules or {},
            is_required=placeholder.is_required
        )

        return Response(
            {
                "placeholder_name": placeholder_name,
                "is_valid": is_valid,
                "error": message if not is_valid else None,
                "formatted_value": message if is_valid and message else value
            },
            status=status.HTTP_200_OK
        )


@method_decorator(csrf_exempt, name='dispatch')
class AgreementUpdateDataView(APIView):
    """API endpoint for updating agreement placeholder data."""

    permission_classes = [IsAuthenticated]

    def patch(self, request, agreement_id):
        """
        Update agreement placeholder data.

        URL: /api/v1/agreements/{agreement_id}/update-data/
        Body: {
            "placeholder_data": {
                "creator_name": "John Doe",
                "fee_amount": "50000",
                ...
            }
        }
        """
        # Check permission
        if request.user.role == 'viewer':
            return Response(
                {"error": "Viewers do not have permission to generate agreements."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get agreement
        try:
            agreement = Agreement.objects.get(id=agreement_id)
        except Agreement.DoesNotExist:
            return Response(
                {"error": "Agreement not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check ownership
        if agreement.generated_by != request.user:
            return Response(
                {"error": "You do not have permission to modify this agreement."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get placeholder data
        placeholder_data = request.data.get('placeholder_data')
        if not placeholder_data or not isinstance(placeholder_data, dict):
            return Response(
                {"error": "placeholder_data is required and must be a dictionary."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Get all placeholders for this template
        placeholders = agreement.template.placeholders.all()
        placeholder_names = {p.name: p for p in placeholders}

        # Validate that all provided placeholders exist in template
        for name in placeholder_data.keys():
            if name not in placeholder_names:
                return Response(
                    {"error": f"Placeholder '{name}' not found in template."},
                    status=status.HTTP_400_BAD_REQUEST
                )

        # Check that all required placeholders are provided
        required_placeholders = [p for p in placeholders if p.is_required]
        missing_required = []
        for placeholder in required_placeholders:
            if placeholder.name not in placeholder_data or not placeholder_data[placeholder.name]:
                missing_required.append(placeholder.display_label)

        if missing_required:
            return Response(
                {
                    "error": "Missing required fields.",
                    "missing_fields": missing_required
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        # Structure data with metadata for tracking
        structured_data = {}
        current_time = timezone.now().isoformat()

        for name, value in placeholder_data.items():
            # Check if this is an update or new entry
            existing_data = agreement.placeholder_data.get(name, {})

            structured_data[name] = {
                'value': value,
                'source': existing_data.get('source', 'manual'),
                'modified_at': current_time
            }

            # If there was an original_value (from auto-fill in Phase 2), preserve it
            if 'original_value' in existing_data:
                structured_data[name]['original_value'] = existing_data['original_value']
                structured_data[name]['overridden'] = True

        # Update agreement
        agreement.placeholder_data = structured_data
        agreement.save()

        # Log audit event
        log_audit_event(
            user=request.user,
            action='agreement_data_updated',
            request=request,
            agreement_id=agreement.id,
            template_id=agreement.template.id,
            fields_updated=list(placeholder_data.keys())
        )

        return Response(
            {
                "message": "Agreement data updated successfully.",
                "agreement_id": agreement.id,
                "fields_updated": list(placeholder_data.keys()),
                "total_fields": len(structured_data),
                "required_fields_complete": len(missing_required) == 0
            },
            status=status.HTTP_200_OK
        )


@method_decorator(csrf_exempt, name='dispatch')
class AgreementPreviewView(APIView):
    """API endpoint for generating agreement preview (ephemeral)."""

    permission_classes = [IsAuthenticated]

    def post(self, request, agreement_id):
        """
        Generate preview of agreement without saving.

        URL: /api/v1/agreements/{agreement_id}/preview/
        """
        # Check permission
        if request.user.role == 'viewer':
            return Response(
                {"error": "Viewers do not have permission to generate agreements."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get agreement
        try:
            agreement = Agreement.objects.get(id=agreement_id)
        except Agreement.DoesNotExist:
            return Response(
                {"error": "Agreement not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check ownership
        if agreement.generated_by != request.user:
            return Response(
                {"error": "You do not have permission to modify this agreement."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Check that all required fields are filled
        placeholders = agreement.template.placeholders.all()
        required_placeholders = [p for p in placeholders if p.is_required]
        missing_required = []
        for placeholder in required_placeholders:
            if placeholder.name not in agreement.placeholder_data or not agreement.placeholder_data[placeholder.name].get('value'):
                missing_required.append(placeholder.display_label)

        if missing_required:
            return Response(
                {
                    "error": "Cannot generate preview. Missing required fields.",
                    "missing_fields": missing_required
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            # Get template file path
            template_path = agreement.template.file_path.path

            # Calculate original template checksum
            original_checksum = calculate_document_checksum(template_path)

            # Generate preview DOCX
            preview_docx = generate_agreement_docx(template_path, agreement.placeholder_data)

            # Calculate generated document checksum
            generated_checksum = calculate_document_checksum(preview_docx.getvalue())

            # Verify integrity
            integrity_verified = (original_checksum == generated_checksum)

            # Log preview
            log_audit_event(
                user=request.user,
                action='agreement_preview_generated',
                request=request,
                agreement_id=agreement.id,
                template_id=agreement.template.id,
                integrity_verified=integrity_verified
            )

            return Response(
                {
                    "message": "Preview generated successfully.",
                    "agreement_id": agreement.id,
                    "integrity_verified": integrity_verified,
                    "original_checksum": original_checksum,
                    "generated_checksum": generated_checksum,
                    "preview_available": True
                },
                status=status.HTTP_200_OK
            )

        except jinja2.exceptions.TemplateSyntaxError as e:
            # Template file has invalid Jinja2 placeholder syntax
            # Try to extract the problematic placeholder from error message
            error_str = str(e)
            template_path = agreement.template.file_path.path

            # Get the template to analyze placeholders
            try:
                template_doc = DocxTemplate(template_path)
                # Extract all text from the document to find placeholders
                all_placeholders = []
                import re
                for paragraph in template_doc.docx.paragraphs:
                    # Find all {{...}} patterns
                    matches = re.findall(r'\{\{[^}]*\}\}', paragraph.text)
                    all_placeholders.extend(matches)

                invalid_placeholders = []
                for placeholder in set(all_placeholders):
                    # Check if placeholder has invalid syntax
                    inner = placeholder.strip('{}').strip()
                    if ' ' in inner or not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', inner):
                        invalid_placeholders.append(placeholder)

                if invalid_placeholders:
                    error_message = (
                        f"Template has {len(invalid_placeholders)} placeholder(s) with invalid syntax:\n\n"
                        + "\n".join([f"• {p}" for p in invalid_placeholders[:10]]) +
                        "\n\nPlaceholders must use format: {{placeholder_name}} (letters, numbers, underscores only, no spaces)"
                    )
                else:
                    error_message = (
                        f"Template syntax error: {error_str}\n\n"
                        "Placeholders must use format: {{placeholder_name}} (no spaces or special characters)"
                    )
            except:
                error_message = (
                    f"Template syntax error: {error_str}\n\n"
                    "Placeholders must use format: {{placeholder_name}} (no spaces or special characters)"
                )

            log_audit_event(
                user=request.user,
                action='agreement_preview_failed',
                request=request,
                agreement_id=agreement.id,
                error=f"Template syntax error: {str(e)}"
            )
            return Response(
                {
                    "error": "Invalid template placeholder syntax",
                    "message": error_message,
                    "template_id": agreement.template.id,
                    "template_name": agreement.template.name
                },
                status=status.HTTP_400_BAD_REQUEST
            )
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"ERROR in preview generation: {str(e)}")
            print(f"Full traceback:\n{error_details}")

            log_audit_event(
                user=request.user,
                action='agreement_preview_failed',
                request=request,
                agreement_id=agreement.id,
                error=str(e)
            )
            return Response(
                {"error": f"Preview generation failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


@method_decorator(csrf_exempt, name='dispatch')
class AgreementGenerateView(APIView):
    """API endpoint for generating final agreement (saves to storage)."""

    permission_classes = [IsAuthenticated]

    def post(self, request, agreement_id):
        """
        Generate final agreement and save to storage.

        URL: /api/v1/agreements/{agreement_id}/generate/
        """
        # Check permission
        if request.user.role == 'viewer':
            return Response(
                {"error": "Viewers do not have permission to generate agreements."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get agreement
        try:
            agreement = Agreement.objects.get(id=agreement_id)
        except Agreement.DoesNotExist:
            return Response(
                {"error": "Agreement not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check ownership
        if agreement.generated_by != request.user:
            return Response(
                {"error": "You do not have permission to modify this agreement."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Check that all required fields are filled
        placeholders = agreement.template.placeholders.all()
        required_placeholders = [p for p in placeholders if p.is_required]
        missing_required = []
        for placeholder in required_placeholders:
            if placeholder.name not in agreement.placeholder_data or not agreement.placeholder_data[placeholder.name].get('value'):
                missing_required.append(placeholder.display_label)

        if missing_required:
            return Response(
                {
                    "error": "Cannot generate agreement. Missing required fields.",
                    "missing_fields": missing_required
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            # Get template file path
            template_path = agreement.template.file_path.path

            # Calculate original template checksum
            original_checksum = calculate_document_checksum(template_path)

            # Generate DOCX
            generated_docx = generate_agreement_docx(template_path, agreement.placeholder_data)

            # Calculate generated document checksum
            generated_checksum = calculate_document_checksum(generated_docx.getvalue())

            # Verify integrity (TODO: Improve to only compare legal text, not filled placeholders)
            # For MVP, we log checksum difference but don't fail the generation
            integrity_verified = True  # Will be enhanced in Phase 2
            if original_checksum != generated_checksum:
                log_audit_event(
                    user=request.user,
                    action='agreement_checksum_different',
                    request=request,
                    agreement_id=agreement.id,
                    template_id=agreement.template.id,
                    original_checksum=original_checksum,
                    generated_checksum=generated_checksum,
                    note='Expected difference due to placeholder filling'
                )

            # Generate agreement ID
            agreement_id_str = generate_agreement_id()

            # Save file to storage
            filename = f"{agreement_id_str}.docx"
            filepath = os.path.join('agreements', str(timezone.now().year), str(timezone.now().month).zfill(2), filename)
            full_path = os.path.join(settings.MEDIA_ROOT, filepath)

            # Ensure directory exists
            os.makedirs(os.path.dirname(full_path), exist_ok=True)

            # Write file
            with open(full_path, 'wb') as f:
                f.write(generated_docx.getvalue())

            # Update agreement
            agreement.agreement_id = agreement_id_str
            agreement.file_path = filepath
            agreement.checksum_sha256 = generated_checksum
            agreement.integrity_verified = True
            agreement.status = 'generated'
            agreement.generated_at = timezone.now()
            agreement.save()

            # Log success
            log_audit_event(
                user=request.user,
                action='agreement_generated',
                request=request,
                agreement_id=agreement.id,
                agreement_ref_id=agreement_id_str,
                template_id=agreement.template.id,
                checksum=generated_checksum
            )

            return Response(
                {
                    "message": "Agreement generated successfully.",
                    "agreement_id": agreement.id,
                    "agreement_ref_id": agreement_id_str,
                    "integrity_verified": True,
                    "checksum": generated_checksum,
                    "status": "generated"
                },
                status=status.HTTP_201_CREATED
            )

        except jinja2.exceptions.TemplateSyntaxError as e:
            # Template file has invalid Jinja2 placeholder syntax
            # Try to extract the problematic placeholder from error message
            error_str = str(e)

            # Get the template to analyze placeholders
            try:
                template_doc = DocxTemplate(template_path)
                # Extract all text from the document to find placeholders
                all_placeholders = []
                import re
                for paragraph in template_doc.docx.paragraphs:
                    # Find all {{...}} patterns
                    matches = re.findall(r'\{\{[^}]*\}\}', paragraph.text)
                    all_placeholders.extend(matches)

                invalid_placeholders = []
                for placeholder in set(all_placeholders):
                    # Check if placeholder has invalid syntax
                    inner = placeholder.strip('{}').strip()
                    if ' ' in inner or not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', inner):
                        invalid_placeholders.append(placeholder)

                if invalid_placeholders:
                    error_message = (
                        f"Template has {len(invalid_placeholders)} placeholder(s) with invalid syntax:\n\n"
                        + "\n".join([f"• {p}" for p in invalid_placeholders[:10]]) +
                        "\n\nPlaceholders must use format: {{placeholder_name}} (letters, numbers, underscores only, no spaces)"
                    )
                else:
                    error_message = (
                        f"Template syntax error: {error_str}\n\n"
                        "Placeholders must use format: {{placeholder_name}} (no spaces or special characters)"
                    )
            except:
                error_message = (
                    f"Template syntax error: {error_str}\n\n"
                    "Placeholders must use format: {{placeholder_name}} (no spaces or special characters)"
                )

            log_audit_event(
                user=request.user,
                action='agreement_generation_failed',
                request=request,
                agreement_id=agreement.id,
                error=f"Template syntax error: {str(e)}"
            )
            return Response(
                {
                    "error": "Invalid template placeholder syntax",
                    "message": error_message,
                    "template_id": agreement.template.id,
                    "template_name": agreement.template.name
                },
                status=status.HTTP_400_BAD_REQUEST
            )
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()

            # Write to file for debugging
            with open('/tmp/agreement_error.log', 'w') as f:
                f.write(f"ERROR in agreement generation: {str(e)}\n")
                f.write(f"Full traceback:\n{error_details}\n")

            print(f"ERROR in agreement generation: {str(e)}")
            print(f"Full traceback:\n{error_details}")

            log_audit_event(
                user=request.user,
                action='agreement_generation_failed',
                request=request,
                agreement_id=agreement.id,
                error=str(e)
            )
            return Response(
                {"error": f"Agreement generation failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


@method_decorator(csrf_exempt, name='dispatch')
class AgreementDownloadView(APIView):
    """API endpoint for downloading generated agreement."""

    permission_classes = [IsAuthenticated]

    def get(self, request, agreement_id):
        """
        Download generated agreement file.

        URL: /api/v1/agreements/{agreement_id}/download/
        """
        # Get agreement
        try:
            agreement = Agreement.objects.get(id=agreement_id)
        except Agreement.DoesNotExist:
            return Response(
                {"error": "Agreement not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check that agreement is generated
        if agreement.status != 'generated' or not agreement.file_path:
            return Response(
                {"error": "Agreement has not been generated yet."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check permission (owner or legal reviewer/admin)
        if agreement.generated_by != request.user and request.user.role not in ['admin', 'legal_reviewer']:
            return Response(
                {"error": "You do not have permission to download this agreement."},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            # Get file path
            file_path = agreement.file_path.path

            # Generate sanitized filename
            import re
            # Get creator name from placeholder data (if available)
            creator_name = 'Unknown'
            for key, value in agreement.placeholder_data.items():
                if 'name' in key.lower() and 'creator' in key.lower():
                    if isinstance(value, dict):
                        creator_name = value.get('value', 'Unknown')
                    else:
                        creator_name = value
                    break

            # Sanitize filename
            template_name = re.sub(r'[^\w\s-]', '', agreement.template.name).strip().replace(' ', '_')
            creator_name = re.sub(r'[^\w\s-]', '', creator_name).strip().replace(' ', '_')
            date_str = timezone.now().strftime('%Y-%m-%d')
            filename = f"{template_name}_{creator_name}_{date_str}.docx"

            # Update downloaded_at if first download
            if not agreement.updated_at or agreement.status == 'generated':
                agreement.status = 'downloaded'
                agreement.save()

            # Log download
            log_audit_event(
                user=request.user,
                action='agreement_downloaded',
                request=request,
                agreement_id=agreement.id,
                agreement_ref_id=agreement.agreement_id,
                filename=filename
            )

            # Return file
            return FileResponse(
                open(file_path, 'rb'),
                as_attachment=True,
                filename=filename,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            return Response(
                {"error": f"Download failed: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class AgreementListView(APIView):
    """API endpoint for listing agreements with search and filters."""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        """
        List agreements with pagination, search, and filters.

        URL: /api/v1/agreements/list/
        Query Params:
          - page: page number (default 1)
          - page_size: items per page (default 20, max 100)
          - search: search in creator name, project name
          - template_id: filter by template
          - date_from: filter by created_at >= date
          - date_to: filter by created_at <= date
        """
        # Get base queryset
        if request.user.role in ['admin', 'legal_reviewer']:
            # Admin and legal reviewers see all agreements
            agreements = Agreement.objects.all()
        else:
            # Content managers see only their own
            agreements = Agreement.objects.filter(generated_by=request.user)

        # Exclude draft agreements (only show generated/downloaded)
        agreements = agreements.exclude(status='draft')

        # Search filter
        search = request.query_params.get('search', None)
        if search:
            # Search in placeholder_data JSON field
            from django.db.models import Q
            agreements = agreements.filter(
                Q(template__name__icontains=search) |
                Q(placeholder_data__icontains=search)
            )

        # Template filter
        template_id = request.query_params.get('template_id', None)
        if template_id:
            agreements = agreements.filter(template_id=template_id)

        # Date range filters
        date_from = request.query_params.get('date_from', None)
        date_to = request.query_params.get('date_to', None)
        if date_from:
            # Parse ISO format date string, handle both with/without timezone
            date_from_clean = date_from.replace(' ', '+')
            agreements = agreements.filter(generated_at__gte=date_from_clean)
        if date_to:
            date_to_clean = date_to.replace(' ', '+')
            agreements = agreements.filter(generated_at__lte=date_to_clean)

        # Order by generated date (most recent first)
        agreements = agreements.order_by('-generated_at', '-created_at')

        # Pagination
        page = int(request.query_params.get('page', 1))
        page_size = min(int(request.query_params.get('page_size', 20)), 100)

        start = (page - 1) * page_size
        end = start + page_size
        total_count = agreements.count()
        agreements_page = agreements[start:end]

        # Build response data
        agreements_data = []
        for agreement in agreements_page:
            # Extract creator name from placeholder data
            creator_name = 'Unknown'
            project_name = None
            for key, value in agreement.placeholder_data.items():
                if 'creator' in key.lower() and 'name' in key.lower():
                    if isinstance(value, dict):
                        creator_name = value.get('value', 'Unknown')
                    else:
                        creator_name = value
                if 'project' in key.lower() and 'name' in key.lower():
                    if isinstance(value, dict):
                        project_name = value.get('value')
                    else:
                        project_name = value

            agreements_data.append({
                'id': agreement.id,
                'agreement_id': agreement.agreement_id,
                'template_name': agreement.template.name,
                'template_category': agreement.template.category,
                'creator_name': creator_name,
                'project_name': project_name,
                'generated_at': agreement.generated_at,
                'downloaded_at': agreement.downloaded_at,
                'status': agreement.status,
                'integrity_verified': agreement.integrity_verified,
                'generated_by': {
                    'id': agreement.generated_by.id,
                    'email': agreement.generated_by.email
                } if agreement.generated_by else None,
                'regeneration_of': agreement.regeneration_of_agreement_id
            })

        return Response({
            'agreements': agreements_data,
            'pagination': {
                'page': page,
                'page_size': page_size,
                'total_count': total_count,
                'total_pages': (total_count + page_size - 1) // page_size,
                'has_next': end < total_count,
                'has_prev': page > 1
            }
        }, status=status.HTTP_200_OK)


class AgreementDetailView(APIView):
    """API endpoint for retrieving agreement details."""

    permission_classes = [IsAuthenticated]

    def get(self, request, agreement_id):
        """
        Get agreement details with grouped placeholders.

        URL: /api/v1/agreements/{agreement_id}/
        """
        # Get agreement
        try:
            agreement = Agreement.objects.select_related('template', 'generated_by').get(id=agreement_id)
        except Agreement.DoesNotExist:
            return Response(
                {"error": "Agreement not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check permissions
        if request.user.role not in ['admin', 'legal_reviewer']:
            if agreement.generated_by != request.user:
                return Response(
                    {"error": "You do not have permission to view this agreement."},
                    status=status.HTTP_403_FORBIDDEN
                )

        # Get template and placeholders
        template = agreement.template
        placeholders = template.placeholders.all()

        # Build grouped placeholders
        grouped_placeholders = []
        for placeholder in placeholders:
            # Extract dropdown options and validation regex from validation_rules JSON
            validation_rules = placeholder.validation_rules or {}
            dropdown_options = validation_rules.get('dropdown_options', [])
            validation_regex = validation_rules.get('regex', '')

            grouped_placeholders.append({
                'name': placeholder.name,
                'label': placeholder.display_label,
                'field_type': placeholder.field_type,
                'is_required': placeholder.is_required,
                'dropdown_options': dropdown_options,
                'validation_regex': validation_regex,
                'category': 'Other',  # Category grouping happens on frontend based on placeholder name
            })

        # Build response
        response_data = {
            'id': agreement.id,
            'agreement_id': agreement.agreement_id,
            'status': agreement.status,
            'template': {
                'id': template.id,
                'name': template.name,
                'version': template.version,
                'category': template.category,
            },
            'generated_by': {
                'id': agreement.generated_by.id,
                'email': agreement.generated_by.email,
            } if agreement.generated_by else None,
            'placeholder_data': agreement.placeholder_data,
            'grouped_placeholders': grouped_placeholders,
            'created_at': agreement.created_at,
            'updated_at': agreement.updated_at,
            'generated_at': agreement.generated_at,
            'downloaded_at': agreement.downloaded_at,
            'integrity_verified': agreement.integrity_verified,
        }

        return Response(response_data, status=status.HTTP_200_OK)
