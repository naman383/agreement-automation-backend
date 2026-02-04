"""Utility functions for extracting and managing placeholders from DOCX templates."""
import re
from docx import Document
from .models import Placeholder


def extract_placeholders_from_docx(file_obj):
    """
    Extract all {{placeholder}} patterns from a DOCX file.
    Returns a list of unique placeholder names.
    """
    try:
        # Reset file pointer
        file_obj.seek(0)
        
        # Load the DOCX
        doc = Document(file_obj)
        
        # Find all placeholders in format {{placeholder_name}}
        placeholder_pattern = r'\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}'
        placeholders = set()
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(placeholder_pattern, paragraph.text)
            placeholders.update(matches)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(placeholder_pattern, paragraph.text)
                        placeholders.update(matches)
        
        return sorted(list(placeholders))
    except Exception as e:
        print(f"Error extracting placeholders: {str(e)}")
        return []


def auto_detect_field_type(placeholder_name):
    """Auto-detect field type based on placeholder name."""
    name_lower = placeholder_name.lower()
    
    # Date fields
    if any(keyword in name_lower for keyword in ['date', 'dob', 'birth', 'expiry', 'valid']):
        return 'date'
    
    # Email fields
    if 'email' in name_lower or 'mail' in name_lower:
        return 'email'
    
    # Phone fields
    if any(keyword in name_lower for keyword in ['phone', 'mobile', 'contact', 'tel']):
        return 'phone'
    
    # Number/Amount fields
    if any(keyword in name_lower for keyword in ['amount', 'price', 'cost', 'fee', 'salary', 'payment', 'sum', 'total']):
        return 'currency'
    
    if any(keyword in name_lower for keyword in ['number', 'count', 'quantity', 'age', 'year']):
        return 'number'
    
    # PAN
    if 'pan' in name_lower:
        return 'pan_number'

    # GST
    if 'gst' in name_lower or 'gstin' in name_lower:
        return 'gst_number'
    
    # Default to text
    return 'text'


def create_placeholders_for_template(template, file_obj):
    """
    Extract placeholders from DOCX and create Placeholder records.
    Returns the count of placeholders created.
    """
    placeholder_names = extract_placeholders_from_docx(file_obj)
    
    if not placeholder_names:
        return 0
    
    created_count = 0
    for name in placeholder_names:
        # Auto-detect field type
        field_type = auto_detect_field_type(name)
        
        # Create a human-readable label from the placeholder name
        display_label = name.replace('_', ' ').title()

        # Create the placeholder
        Placeholder.objects.create(
            template=template,
            name=name,
            display_label=display_label,
            field_type=field_type,
            is_required=True,  # Default to required
            position_index=created_count
        )
        created_count += 1
    
    return created_count
