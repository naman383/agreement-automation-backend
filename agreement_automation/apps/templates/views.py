"""Views for template management."""

import hashlib
import zipfile
import io
from datetime import datetime
from rest_framework import status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.authentication import SessionAuthentication
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.http import FileResponse
from django.utils import timezone
from django.db import models as django_models
from docxtpl import DocxTemplate
import jinja2.exceptions


class CsrfExemptSessionAuthentication(SessionAuthentication):
    """Session authentication without CSRF enforcement."""
    def enforce_csrf(self, request):
        return  # Do not enforce CSRF

from .models import Template
from .serializers import TemplateListSerializer, TemplateUploadSerializer
from agreement_automation.apps.audit.utils import log_audit_event


def contains_macros(file_obj):
    """Check if DOCX file contains VBA macros."""
    try:
        file_obj.seek(0)  # Reset file pointer
        with zipfile.ZipFile(file_obj, 'r') as zip_ref:
            # Check for vbaProject.bin (contains VBA macros)
            file_list = zip_ref.namelist()
            if 'word/vbaProject.bin' in file_list:
                return True
        file_obj.seek(0)  # Reset file pointer again
        return False
    except Exception:
        return False


def calculate_checksum(file_obj):
    """Calculate SHA-256 checksum of file."""
    file_obj.seek(0)  # Reset file pointer
    sha256_hash = hashlib.sha256()
    for byte_block in iter(lambda: file_obj.read(4096), b""):
        sha256_hash.update(byte_block)
    file_obj.seek(0)  # Reset file pointer
    return sha256_hash.hexdigest()


def validate_docx_structure(file_obj):
    """Validate DOCX file structure by attempting to parse it."""
    try:
        file_obj.seek(0)
        # Try to open and parse the DOCX file
        from docx import Document
        Document(file_obj)
        file_obj.seek(0)
        return True
    except Exception as e:
        # Any exception means the file is corrupt or invalid
        return False



from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
import json

@method_decorator(csrf_exempt, name='dispatch')
class DiagnosticUploadView(APIView):
    """Diagnostic endpoint that captures data AND processes upload."""
    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [AllowAny]  # No auth to test

    def post(self, request):
        """Capture everything AND process the upload."""

        diagnostic_data = {
            "headers": dict(request.headers),
            "cookies": dict(request.COOKIES),
            "method": request.method,
            "content_type": request.content_type,
            "user": str(request.user),
            "is_authenticated": request.user.is_authenticated,
            "data_keys": list(request.data.keys()),
            "files_keys": list(request.FILES.keys()),
            "meta_selected": {
                "REMOTE_ADDR": request.META.get("REMOTE_ADDR"),
                "HTTP_ORIGIN": request.META.get("HTTP_ORIGIN"),
                "HTTP_REFERER": request.META.get("HTTP_REFERER"),
            }
        }

        # Get file details if present
        for key in request.FILES.keys():
            file_obj = request.FILES[key]
            diagnostic_data[f"file_{key}"] = {
                "name": file_obj.name,
                "size": file_obj.size,
                "content_type": file_obj.content_type,
            }

        # Get form data
        for key in request.data.keys():
            if key not in request.FILES:
                diagnostic_data[f"field_{key}"] = str(request.data[key])

        # Write to file for guaranteed capture
        with open('/tmp/diagnostic_capture.json', 'w') as f:
            f.write("=" * 80 + "\n")
            f.write("🔍 DIAGNOSTIC CAPTURE FROM BROWSER:\n")
            f.write(json.dumps(diagnostic_data, indent=2))
            f.write("\n" + "=" * 80 + "\n")

        # Now try to actually process the upload
        try:
            serializer = TemplateUploadSerializer(data=request.data)

            if not serializer.is_valid():
                with open('/tmp/diagnostic_capture.json', 'a') as f:
                    f.write("❌ VALIDATION FAILED:\n")
                    f.write(json.dumps(serializer.errors, indent=2) + "\n")
                return Response({
                    "status": "validation_failed",
                    "errors": serializer.errors,
                    "diagnostic": diagnostic_data
                }, status=status.HTTP_400_BAD_REQUEST)

            with open('/tmp/diagnostic_capture.json', 'a') as f:
                f.write("✅ VALIDATION PASSED\n")

            file = serializer.validated_data['file']
            name = serializer.validated_data['name']
            category = serializer.validated_data['category']

            # Skip macro check for diagnostic
            # Skip structure check for diagnostic

            # Calculate checksum
            checksum = calculate_checksum(file)

            # Save file
            file_path = default_storage.save(
                f'templates/{file.name}',
                ContentFile(file.read())
            )

            # Create template (use AnonymousUser for diagnostic)
            from django.contrib.auth import get_user_model
            User = get_user_model()
            test_user = User.objects.filter(email='qa-test@example.com').first()

            template = Template.objects.create(
                name=name,
                category=category,
                file_path=file_path,
                checksum_sha256=checksum,
                status='draft',
                uploaded_by=test_user if test_user else None
            )

            with open('/tmp/diagnostic_capture.json', 'a') as f:
                f.write(f"✅ TEMPLATE CREATED: ID={template.id}\n")

            # Extract placeholders from DOCX
            from .placeholder_utils import create_placeholders_for_template
            file.seek(0)  # Reset file pointer
            placeholder_count = create_placeholders_for_template(template, file)

            with open('/tmp/diagnostic_capture.json', 'a') as f:
                f.write(f"✅ EXTRACTED {placeholder_count} PLACEHOLDERS\n")

            return Response({
                "status": "success",
                "id": template.id,
                "message": "Diagnostic upload successful",
                "diagnostic": diagnostic_data
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            with open('/tmp/diagnostic_capture.json', 'a') as f:
                f.write(f"❌ EXCEPTION: {str(e)}\n")
                import traceback
                f.write(traceback.format_exc())
            return Response({
                "status": "exception",
                "error": str(e),
                "diagnostic": diagnostic_data
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class TemplateListView(APIView):
    """API endpoint for listing templates."""

    permission_classes = [IsAuthenticated]

    def get(self, request):
        """
        List templates based on user role.

        URL: /api/v1/templates/
        Query Params:
          - status: draft/approved/deprecated/all (default: based on role)
          - category: filter by category

        RBAC:
          - Admin: Can view all templates (all statuses)
          - Legal Reviewer: Can view all templates (all statuses)
          - Content Manager: Can only view approved templates
          - Viewer: Can only view approved templates
        """
        # RBAC filtering
        if request.user.is_admin or request.user.role == 'legal_reviewer':
            # Admin and Legal Reviewer can view all templates
            templates = Template.objects.all()

            # Optional status filter
            status_filter = request.query_params.get('status', None)
            if status_filter and status_filter != 'all':
                templates = templates.filter(status=status_filter)
        else:
            # Content Manager, Viewer, and other roles can only view approved templates
            templates = Template.objects.filter(status='approved')

        # Optional category filter
        category_filter = request.query_params.get('category', None)
        if category_filter:
            templates = templates.filter(category=category_filter)

        serializer = TemplateListSerializer(templates, many=True)

        return Response(
            {"templates": serializer.data},
            status=status.HTTP_200_OK
        )


class TemplateDetailView(APIView):
    """API endpoint for getting a single template's details."""

    permission_classes = [IsAuthenticated]

    def get(self, request, template_id):
        """
        Get template details by ID.

        URL: /api/v1/templates/<template_id>/
        """
        try:
            template = Template.objects.get(id=template_id)

            # RBAC check
            if not (request.user.is_admin or request.user.role == 'legal_reviewer'):
                # Non-admin users can only view approved templates
                if template.status != 'approved':
                    return Response(
                        {"error": "You do not have permission to view this template."},
                        status=status.HTTP_403_FORBIDDEN
                    )

            from .serializers import TemplateDetailSerializer
            serializer = TemplateDetailSerializer(template)
            return Response(serializer.data, status=status.HTTP_200_OK)

        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )


@method_decorator(csrf_exempt, name='dispatch')
class PlaceholderConfigView(APIView):
    """API endpoint for configuring template placeholders."""

    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request, template_id):
        """
        Update placeholder configurations for a template.

        URL: /api/v1/templates/<template_id>/placeholders/
        """
        try:
            from .models import Placeholder

            template = Template.objects.get(id=template_id)

            # Check permissions
            if not request.user.is_admin:
                return Response(
                    {"error": "You do not have permission to configure placeholders."},
                    status=status.HTTP_403_FORBIDDEN
                )

            placeholders_data = request.data.get('placeholders', [])

            if not isinstance(placeholders_data, list):
                return Response(
                    {"error": "Placeholders must be a list."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Update or create placeholders
            for placeholder_data in placeholders_data:
                placeholder_id = placeholder_data.get('id')

                if placeholder_id:
                    # Update existing placeholder
                    try:
                        placeholder = Placeholder.objects.get(id=placeholder_id, template=template)
                        placeholder.display_label = placeholder_data.get('label', placeholder.display_label)
                        placeholder.field_type = placeholder_data.get('field_type', placeholder.field_type)
                        placeholder.is_required = placeholder_data.get('is_required', placeholder.is_required)
                        placeholder.position_index = placeholder_data.get('position_index', placeholder.position_index)
                        placeholder.validation_rules = placeholder_data.get('validation_rules', placeholder.validation_rules)
                        placeholder.save()
                    except Placeholder.DoesNotExist:
                        return Response(
                            {"error": f"Placeholder with ID {placeholder_id} not found."},
                            status=status.HTTP_404_NOT_FOUND
                        )
                else:
                    # Create new placeholder
                    Placeholder.objects.create(
                        template=template,
                        name=placeholder_data.get('name', ''),
                        display_label=placeholder_data.get('label', ''),
                        field_type=placeholder_data.get('field_type', 'text'),
                        is_required=placeholder_data.get('is_required', True),
                        position_index=placeholder_data.get('position_index', 0),
                        validation_rules=placeholder_data.get('validation_rules', {})
                    )

            # Log audit event
            log_audit_event(
                user=request.user,
                action='template_placeholder_config',
                details={'template_id': template.id, 'template_name': template.name}
            )

            return Response(
                {"message": "Placeholders configured successfully."},
                status=status.HTTP_200_OK
            )

        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            print(f"ERROR configuring placeholders: {str(e)}")
            return Response(
                {"error": f"Failed to configure placeholders: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


@method_decorator(csrf_exempt, name='dispatch')
class TemplateUploadView(APIView):
    """API endpoint for uploading templates."""

    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Upload template file.

        URL: /api/v1/templates/upload/
        """
        # Check if user is admin
        if not request.user.is_admin:
            return Response(
                {"error": "You do not have permission to perform this action."},
                status=status.HTTP_403_FORBIDDEN
            )

        # DEBUG: Log request data
        print("UPLOAD DEBUG - Request data keys:", request.data.keys())
        print("UPLOAD DEBUG - Request FILES:", request.FILES.keys())
        for key in request.data.keys():
            if key in request.FILES:
                file_obj = request.FILES[key]
                print(f"UPLOAD DEBUG - File '{key}': name={file_obj.name}, size={file_obj.size}, content_type={file_obj.content_type}")
            else:
                print(f"UPLOAD DEBUG - Field '{key}': {request.data[key]}")
        
        serializer = TemplateUploadSerializer(data=request.data)

        if not serializer.is_valid():
            print("UPLOAD DEBUG - Validation FAILED:", serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        print("UPLOAD DEBUG - Validation PASSED")

        file = serializer.validated_data['file']
        name = serializer.validated_data['name']
        category = serializer.validated_data['category']

        # Check for macros
        if contains_macros(file):
            # Log security event
            log_audit_event(
                user=request.user,
                action='template_upload_rejected_macros',
                request=request,
                template_name=name,
                reason='File contains macros'
            )
            return Response(
                {"error": "File contains macros. Macro-enabled documents are not allowed for security reasons."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Validate DOCX structure
        if not validate_docx_structure(file):
            return Response(
                {"error": "Invalid DOCX file. File may be corrupted or not a valid DOCX document."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Calculate checksum
        checksum = calculate_checksum(file)

        # Save file to storage
        file_path = default_storage.save(
            f'templates/{file.name}',
            ContentFile(file.read())
        )

        # Create template record
        template = Template.objects.create(
            name=name,
            category=category,
            file_path=file_path,
            checksum_sha256=checksum,
            status='draft',
            uploaded_by=request.user
        )

        # Extract placeholders from DOCX
        from .placeholder_utils import create_placeholders_for_template
        file.seek(0)  # Reset file pointer
        placeholder_count = create_placeholders_for_template(template, file)
        print(f"UPLOAD DEBUG - Extracted {placeholder_count} placeholders")

        # Log upload event
        log_audit_event(
            user=request.user,
            action='template_uploaded',
            request=request,
            template_name=name,
            template_id=template.id,
            checksum=checksum
        )

        return Response(
            {
                "message": "Template uploaded successfully. Status: Draft",
                "template": {
                    "id": template.id,
                    "name": template.name,
                    "category": template.category,
                    "status": template.status,
                    "checksum": template.checksum_sha256,
                    "created_at": template.created_at
                }
            },
            status=status.HTTP_201_CREATED
        )


def generate_dummy_data(placeholders):
    """Generate dummy data for placeholders based on field type."""
    dummy_data = {}

    for placeholder in placeholders:
        field_type = placeholder.field_type
        name = placeholder.name

        # Special handling for common placeholder names
        if 'name' in name.lower() and 'creator' in name.lower():
            dummy_data[name] = 'John Doe'
        elif 'name' in name.lower():
            dummy_data[name] = 'Sample Name'
        elif field_type == 'text':
            dummy_data[name] = 'Sample Text'
        elif field_type == 'number':
            dummy_data[name] = '12345'
        elif field_type == 'date':
            dummy_data[name] = datetime.now().strftime('%Y-%m-%d')
        elif field_type == 'currency':
            dummy_data[name] = '50,000'
        elif field_type == 'pan_number':
            dummy_data[name] = 'ABCDE1234F'
        elif field_type == 'gst_number':
            dummy_data[name] = '22ABCDE1234F1Z5'
        elif field_type == 'dropdown':
            # Check if validation_rules has options
            if placeholder.validation_rules and 'options' in placeholder.validation_rules:
                dummy_data[name] = placeholder.validation_rules['options'][0]
            else:
                dummy_data[name] = 'Sample Option'
        elif field_type == 'checkbox':
            dummy_data[name] = 'Yes'
        else:
            dummy_data[name] = 'Sample Value'

    return dummy_data


@method_decorator(csrf_exempt, name='dispatch')
class TemplatePreviewView(APIView):
    """API endpoint for previewing templates with dummy data."""

    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request, template_id):
        """
        Generate preview of template with dummy or custom data.

        URL: /api/v1/templates/<id>/preview/
        """
        # Check if user is admin
        if not request.user.is_admin:
            return Response(
                {"error": "You do not have permission to perform this action."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get template
        try:
            template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Get placeholders for this template
        placeholders = template.placeholders.all()

        if not placeholders.exists():
            return Response(
                {"error": "No placeholders found for this template. Please tag placeholders first."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Use custom data if provided, otherwise generate dummy data
        if request.data and 'placeholder_data' in request.data:
            context_data = request.data['placeholder_data']
        else:
            context_data = generate_dummy_data(placeholders)

        try:
            # Load template file
            template_file_path = template.file_path.path
            doc = DocxTemplate(template_file_path)

            # Render template with data
            doc.render(context_data)

            # Save to BytesIO buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{template.name}_preview_{timestamp}.docx"

            # Log preview event
            log_audit_event(
                user=request.user,
                action='template_preview_generated',
                details={
                    'template_id': template.id,
                    'template_name': template.name
                }
            )

            # Return file response
            response = FileResponse(
                buffer,
                as_attachment=True,
                filename=filename,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            return response

        except jinja2.exceptions.TemplateSyntaxError as e:
            # Template file has invalid Jinja2 placeholder syntax
            error_message = (
                "Template file contains invalid placeholder syntax. "
                "Please ensure all placeholders use the correct format: {{placeholder_name}} "
                "with no spaces or special characters inside the braces."
            )
            print(f"Template syntax error in template {template.id}: {str(e)}")
            return Response(
                {"error": error_message, "details": str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"ERROR generating preview: {str(e)}")
            print(f"Full traceback:\n{error_details}")
            return Response(
                {"error": f"Failed to generate preview: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


@method_decorator(csrf_exempt, name='dispatch')
class TemplateApproveView(APIView):
    """API endpoint for approving templates."""

    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request, template_id):
        """
        Approve template for production use.

        URL: /api/v1/templates/<id>/approve/
        """
        # Check if user is admin or legal_reviewer
        if not (request.user.is_admin or request.user.role == 'legal_reviewer'):
            return Response(
                {"error": "Access Denied. Only admins and legal reviewers can approve templates."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get template
        try:
            template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check if already approved
        if template.status == 'approved':
            return Response(
                {"error": "Template is already approved."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check if template is deprecated
        if template.status == 'deprecated':
            return Response(
                {"error": "Cannot approve a deprecated template."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check if template has placeholders
        if not template.placeholders.exists():
            return Response(
                {"error": "Template must have at least one placeholder before approval."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Approve template
        template.status = 'approved'
        template.approved_by = request.user
        template.approved_at = timezone.now()
        template.save()

        # Auto-deprecate old approved versions in the same version chain
        if template.parent_template:
            # This is a new version - deprecate the parent and all other approved versions
            root_template = template.parent_template
            Template.objects.filter(
                django_models.Q(id=root_template.id) | django_models.Q(parent_template=root_template),
                status='approved'
            ).exclude(id=template.id).update(status='deprecated')
        else:
            # This is the root template - deprecate any other approved versions
            Template.objects.filter(
                parent_template=template,
                status='approved'
            ).update(status='deprecated')

        # Log approval event
        log_audit_event(
            user=request.user,
            action='template_approved',
            request=request,
            template_id=template.id,
            template_name=template.name,
            checksum=template.checksum_sha256
        )

        return Response(
            {
                "message": "Template approved. Content managers can now use this template.",
                "template": {
                    "id": template.id,
                    "name": template.name,
                    "status": template.status,
                    "approved_by": request.user.email,
                    "approved_at": template.approved_at
                }
            },
            status=status.HTTP_200_OK
        )


@method_decorator(csrf_exempt, name='dispatch')
class TemplateNewVersionView(APIView):
    """API endpoint for creating new template versions."""

    permission_classes = [IsAuthenticated]

    def post(self, request, template_id):
        """
        Create new version of existing template.

        URL: /api/v1/templates/<id>/new-version/
        """
        # Check if user is admin
        if not request.user.is_admin:
            return Response(
                {"error": "You do not have permission to perform this action."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get parent template
        try:
            parent_template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Parent template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check if parent template is approved
        if parent_template.status != 'approved':
            return Response(
                {"error": "Can only create new version from approved templates."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Validate file upload
        if 'file' not in request.data:
            return Response(
                {"error": "File is required."},
                status=status.HTTP_400_BAD_REQUEST
            )

        file = request.data['file']

        # Validate file type
        if not file.name.lower().endswith('.docx'):
            return Response(
                {"error": "Only DOCX format is supported."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Validate file size
        if file.size > 10 * 1024 * 1024:  # 10MB
            return Response(
                {"error": "File size exceeds 10MB limit."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check for macros
        if contains_macros(file):
            # Log security event
            log_audit_event(
                user=request.user,
                action='template_version_upload_rejected_macros',
                request=request,
                template_name=parent_template.name,
                reason='File contains macros'
            )
            return Response(
                {"error": "File contains macros. Macro-enabled documents are not allowed for security reasons."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Validate DOCX structure
        if not validate_docx_structure(file):
            return Response(
                {"error": "Invalid DOCX file. File may be corrupted or not a valid DOCX document."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Calculate checksum
        checksum = calculate_checksum(file)

        # Find the highest version number in the version chain
        # Start from the parent template
        root_template = parent_template.parent_template or parent_template
        all_versions = Template.objects.filter(
            django_models.Q(id=root_template.id) | django_models.Q(parent_template=root_template)
        )
        max_version = all_versions.aggregate(django_models.Max('version'))['version__max'] or 0
        new_version_number = max_version + 1

        # Save file to storage
        file_path = default_storage.save(
            f'templates/{file.name}',
            ContentFile(file.read())
        )

        # Create new template version
        new_template = Template.objects.create(
            name=parent_template.name,
            category=parent_template.category,
            file_path=file_path,
            checksum_sha256=checksum,
            version=new_version_number,
            status='draft',
            parent_template=root_template,
            uploaded_by=request.user
        )

        # Log version upload event
        log_audit_event(
            user=request.user,
            action='template_version_uploaded',
            request=request,
            template_id=new_template.id,
            template_name=new_template.name,
            version=new_version_number,
            parent_template_id=root_template.id,
            checksum=checksum
        )

        return Response(
            {
                "message": f"Template version {new_version_number} uploaded. Status: Draft",
                "template": {
                    "id": new_template.id,
                    "name": new_template.name,
                    "version": new_template.version,
                    "status": new_template.status,
                    "parent_template_id": root_template.id,
                    "checksum": new_template.checksum_sha256,
                    "created_at": new_template.created_at
                }
            },
            status=status.HTTP_201_CREATED
        )


class TemplateVersionListView(APIView):
    """API endpoint for listing template versions."""

    permission_classes = [IsAuthenticated]

    def get(self, request, template_id):
        """
        List all versions of a template.

        URL: /api/v1/templates/<id>/versions/
        """
        # Check if user is admin or legal_reviewer
        if not (request.user.is_admin or request.user.role == 'legal_reviewer'):
            return Response(
                {"error": "You do not have permission to perform this action."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get template
        try:
            template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Find root template (could be the template itself or its parent)
        root_template = template.parent_template or template

        # Get all versions in the chain
        versions = Template.objects.filter(
            django_models.Q(id=root_template.id) | django_models.Q(parent_template=root_template)
        ).order_by('-version')

        # Serialize versions
        versions_data = []
        for version in versions:
            versions_data.append({
                'id': version.id,
                'name': version.name,
                'version': version.version,
                'status': version.status,
                'uploaded_by': version.uploaded_by.email if version.uploaded_by else None,
                'approved_by': version.approved_by.email if version.approved_by else None,
                'approved_at': version.approved_at,
                'created_at': version.created_at,
                'checksum': version.checksum_sha256
            })

        return Response(
            {
                "template_name": root_template.name,
                "versions": versions_data
            },
            status=status.HTTP_200_OK
        )


@method_decorator(csrf_exempt, name='dispatch')
class TemplateDeprecateView(APIView):
    """API endpoint for deprecating templates."""

    permission_classes = [IsAuthenticated]

    def post(self, request, template_id):
        """
        Deprecate template.

        URL: /api/v1/templates/<id>/deprecate/
        """
        # Check if user is admin
        if not request.user.is_admin:
            return Response(
                {"error": "You do not have permission to perform this action."},
                status=status.HTTP_403_FORBIDDEN
            )

        # Get template
        try:
            template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check if already deprecated
        if template.status == 'deprecated':
            return Response(
                {"error": "Template is already deprecated."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Check if template is draft (can only deprecate approved templates)
        if template.status == 'draft':
            return Response(
                {"error": "Cannot deprecate draft templates. Only approved templates can be deprecated."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Get deprecation reason (optional)
        deprecation_reason = request.data.get('reason', 'Manually deprecated by admin')

        # Deprecate template
        template.status = 'deprecated'
        template.deprecated_by = request.user
        template.deprecated_at = timezone.now()
        template.deprecation_reason = deprecation_reason
        template.save()

        # Log deprecation event
        log_audit_event(
            user=request.user,
            action='template_deprecated',
            request=request,
            template_id=template.id,
            template_name=template.name,
            version=template.version,
            reason=deprecation_reason
        )

        return Response(
            {
                "message": f"Template version {template.version} deprecated",
                "template": {
                    "id": template.id,
                    "name": template.name,
                    "version": template.version,
                    "status": template.status,
                    "deprecated_by": request.user.email,
                    "deprecated_at": template.deprecated_at,
                    "deprecation_reason": deprecation_reason
                }
            },
            status=status.HTTP_200_OK
        )


@method_decorator(csrf_exempt, name='dispatch')
class PlaceholderScanView(APIView):
    """API endpoint for scanning and validating placeholders in template."""

    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request, template_id):
        """
        Scan template DOCX file for all placeholders and validate syntax.

        Returns list of placeholders with validation status and auto-fix suggestions.
        """
        # Get template
        try:
            template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check permissions
        if request.user.role == 'viewer':
            return Response(
                {"error": "Insufficient permissions."},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            # Load template document
            doc = DocxTemplate(template.file_path.path)

            # Extract all text and find placeholder patterns
            import re
            all_placeholders = []

            # Scan paragraphs
            for paragraph in doc.docx.paragraphs:
                matches = re.findall(r'\{\{[^}]*\}\}', paragraph.text)
                all_placeholders.extend(matches)

            # Scan tables
            for table in doc.docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(r'\{\{[^}]*\}\}', paragraph.text)
                            all_placeholders.extend(matches)

            # Validate each placeholder
            placeholder_regex = re.compile(r'^[a-zA-Z_][a-zA-Z0-9_]*$')
            results = []
            seen = set()

            for placeholder_full in all_placeholders:
                if placeholder_full in seen:
                    continue
                seen.add(placeholder_full)

                # Extract inner content
                inner = placeholder_full.strip('{}').strip()

                # Check if valid
                is_valid = bool(placeholder_regex.match(inner))

                # Generate auto-fix suggestion if invalid
                auto_fix = None
                issues = []

                if not is_valid:
                    # Generate fixed version
                    fixed = inner

                    # Replace spaces with underscores
                    if ' ' in fixed:
                        fixed = fixed.replace(' ', '_')
                        issues.append('Contains spaces')

                    # Replace hyphens with underscores
                    if '-' in fixed:
                        fixed = fixed.replace('-', '_')
                        issues.append('Contains hyphens')

                    # Remove other special characters
                    fixed = re.sub(r'[^a-zA-Z0-9_]', '', fixed)

                    # Ensure starts with letter or underscore
                    if fixed and not fixed[0].isalpha() and fixed[0] != '_':
                        fixed = '_' + fixed
                        issues.append('Must start with letter or underscore')

                    if not fixed:
                        fixed = 'placeholder'
                        issues.append('Invalid characters removed')

                    auto_fix = f'{{{{{fixed}}}}}'

                results.append({
                    'original': placeholder_full,
                    'inner_text': inner,
                    'is_valid': is_valid,
                    'issues': issues if not is_valid else [],
                    'suggested_fix': auto_fix,
                    'line_preview': placeholder_full
                })

            # Count valid/invalid
            valid_count = sum(1 for r in results if r['is_valid'])
            invalid_count = len(results) - valid_count

            return Response({
                'template_id': template.id,
                'template_name': template.name,
                'total_placeholders': len(results),
                'valid_count': valid_count,
                'invalid_count': invalid_count,
                'placeholders': results,
                'can_fix_automatically': invalid_count > 0
            }, status=status.HTTP_200_OK)

        except Exception as e:
            import traceback
            print(f"Error scanning template: {str(e)}")
            print(traceback.format_exc())
            return Response(
                {"error": f"Failed to scan template: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


@method_decorator(csrf_exempt, name='dispatch')
class VisualEditorView(APIView):
    """API endpoint for visual template editor."""

    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request, template_id):
        """
        Load template for visual editor.
        Returns rendered HTML and metadata.
        """
        try:
            template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Check permissions
        if request.user.role == 'viewer':
            return Response(
                {"error": "Insufficient permissions."},
                status=status.HTTP_403_FORBIDDEN
            )

        try:
            from .services import DocumentRenderService
            from .models import TemplateVisualMapping, PlaceholderVisualMarker

            # Check if we have cached rendering
            try:
                visual_mapping = template.visual_mapping
                # Return cached data
                markers = PlaceholderVisualMarker.objects.filter(
                    mapping=visual_mapping
                ).select_related('placeholder')

                markers_data = [{
                    'id': marker.id,
                    'placeholder_name': marker.placeholder.name,
                    'placeholder_label': marker.placeholder.display_label,
                    'page_number': marker.page_number,
                    'x': marker.x_position,
                    'y': marker.y_position,
                    'width': marker.width,
                    'height': marker.height,
                    'color': marker.marker_color
                } for marker in markers]

                return Response({
                    'template_id': template.id,
                    'template_name': template.name,
                    'html': visual_mapping.rendered_html,
                    'page_data': visual_mapping.page_data,
                    'markers': markers_data,
                    'cached': True
                }, status=status.HTTP_200_OK)

            except TemplateVisualMapping.DoesNotExist:
                # Render for the first time
                render_result = DocumentRenderService.render_docx_to_html(
                    template.file_path.path
                )

                # Create visual mapping cache
                visual_mapping = TemplateVisualMapping.objects.create(
                    template=template,
                    rendered_html=render_result['html'],
                    page_data={
                        'page_count': render_result['page_count'],
                        'placeholders': render_result['placeholders']
                    }
                )

                return Response({
                    'template_id': template.id,
                    'template_name': template.name,
                    'html': render_result['html'],
                    'page_data': visual_mapping.page_data,
                    'markers': [],
                    'cached': False
                }, status=status.HTTP_200_OK)

        except Exception as e:
            import traceback
            print(f"Visual editor error: {str(e)}")
            print(traceback.format_exc())
            return Response(
                {"error": f"Failed to load visual editor: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def post(self, request, template_id):
        """
        Save placeholder markers in bulk.
        Body: { markers: [{placeholder, x, y, width, height, page, fieldType, isRequired, displayLabel, ...}] }
        """
        try:
            template = Template.objects.get(id=template_id)
        except Template.DoesNotExist:
            return Response(
                {"error": "Template not found."},
                status=status.HTTP_404_NOT_FOUND
            )

        try:
            from .models import TemplateVisualMapping, PlaceholderVisualMarker, Placeholder

            # Get or create visual mapping
            visual_mapping, _ = TemplateVisualMapping.objects.get_or_create(
                template=template
            )

            markers_data = request.data.get('markers', [])
            saved_markers = []
            errors = []

            for marker_data in markers_data:
                try:
                    placeholder_name = marker_data.get('placeholder')
                    if not placeholder_name:
                        errors.append({'marker': marker_data, 'error': 'Missing placeholder name'})
                        continue

                    # Get or create placeholder
                    placeholder, _ = Placeholder.objects.get_or_create(
                        template=template,
                        name=placeholder_name,
                        defaults={
                            'display_label': marker_data.get('displayLabel', placeholder_name),
                            'field_type': marker_data.get('fieldType', 'text'),
                            'is_required': marker_data.get('isRequired', True),
                            'position_index': 0
                        }
                    )

                    # Update placeholder configuration
                    placeholder.field_type = marker_data.get('fieldType', placeholder.field_type)
                    placeholder.is_required = marker_data.get('isRequired', placeholder.is_required)
                    placeholder.display_label = marker_data.get('displayLabel', placeholder.display_label)

                    # Update validation rules
                    validation_rules = {}
                    if marker_data.get('validationRegex'):
                        validation_rules['regex'] = marker_data.get('validationRegex')
                    if marker_data.get('dropdownOptions'):
                        validation_rules['dropdown_options'] = marker_data.get('dropdownOptions')

                    placeholder.validation_rules = validation_rules
                    placeholder.save()

                    # Create or update visual marker
                    marker, created = PlaceholderVisualMarker.objects.update_or_create(
                        placeholder=placeholder,
                        defaults={
                            'mapping': visual_mapping,
                            'page_number': marker_data.get('page', 1),
                            'x_position': marker_data.get('x', 0),
                            'y_position': marker_data.get('y', 0),
                            'width': marker_data.get('width', 10),
                            'height': marker_data.get('height', 3),
                            'marker_color': marker_data.get('color', '#DC2626')
                        }
                    )

                    saved_markers.append({
                        'placeholder': placeholder_name,
                        'marker_id': marker.id,
                        'created': created
                    })

                except Exception as e:
                    errors.append({
                        'marker': marker_data,
                        'error': str(e)
                    })

            return Response({
                'message': f'Saved {len(saved_markers)} markers',
                'saved': saved_markers,
                'errors': errors if errors else None
            }, status=status.HTTP_200_OK)

        except Exception as e:
            import traceback
            print(f"Save markers error: {str(e)}")
            print(traceback.format_exc())
            return Response(
                {"error": f"Failed to save markers: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
