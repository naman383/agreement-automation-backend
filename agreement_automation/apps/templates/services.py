"""Services for template document processing and rendering."""

import re
from io import BytesIO
from docxtpl import DocxTemplate
from docx import Document
import mammoth


class DocumentRenderService:
    """Service for rendering DOCX templates to HTML for visual editor."""

    @staticmethod
    def render_docx_to_html(docx_path):
        """
        Convert DOCX file to HTML for display in visual editor.

        Returns dict with:
        - html: Rendered HTML content
        - page_count: Estimated page count
        - placeholders: List of detected placeholders
        """
        try:
            # Use mammoth to convert DOCX to HTML
            with open(docx_path, 'rb') as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                messages = result.messages

            # Extract placeholders from HTML
            placeholders = DocumentRenderService._extract_placeholders(html_content)

            # Estimate page count (rough estimation based on content length)
            # Average page ~3000 characters
            char_count = len(html_content)
            estimated_pages = max(1, (char_count // 3000) + 1)

            return {
                'html': html_content,
                'page_count': estimated_pages,
                'placeholders': placeholders,
                'conversion_messages': [str(m) for m in messages]
            }

        except Exception as e:
            raise Exception(f"Failed to render document: {str(e)}")

    @staticmethod
    def _extract_placeholders(html_content):
        """Extract all {{...}} placeholder patterns from HTML."""
        # Find all {{...}} patterns
        pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(pattern, html_content)

        # Return unique placeholders
        unique_placeholders = list(set(matches))
        return unique_placeholders

    @staticmethod
    def get_document_structure(docx_path):
        """
        Analyze DOCX structure for advanced features.

        Returns information about paragraphs, tables, sections.
        """
        doc = Document(docx_path)

        structure = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections),
            'paragraphs': [],
            'tables': []
        }

        # Analyze paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure['paragraphs'].append({
                'index': i,
                'text': para.text[:100],  # First 100 chars
                'style': para.style.name,
                'has_placeholder': '{{' in para.text
            })

        # Analyze tables
        for i, table in enumerate(doc.tables):
            structure['tables'].append({
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.columns)
            })

        return structure
