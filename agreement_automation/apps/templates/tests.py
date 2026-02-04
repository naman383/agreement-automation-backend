"""Tests for template management."""

import io
import zipfile
from datetime import datetime
from django.test import TestCase
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.utils import timezone
from rest_framework.test import APITestCase, APIClient
from rest_framework import status
from django.urls import reverse

from .models import Template, Placeholder
from agreement_automation.apps.audit.models import AuditLog
from docxtpl import DocxTemplate

User = get_user_model()


class TemplateModelTests(TestCase):
    """Test cases for Template model."""

    def setUp(self):
        """Create test users."""
        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

    def test_template_creation(self):
        """Test creating a template."""
        template = Template.objects.create(
            name='Test Template',
            category='Test Category',
            file_path='templates/test.docx',
            checksum_sha256='abc123',
            uploaded_by=self.admin
        )

        self.assertEqual(template.name, 'Test Template')
        self.assertEqual(template.status, 'draft')
        self.assertEqual(template.version, 1)

    def test_template_string_representation(self):
        """Test template string representation."""
        template = Template.objects.create(
            name='Test Template',
            category='Test Category',
            file_path='templates/test.docx',
            checksum_sha256='abc123',
            uploaded_by=self.admin
        )

        self.assertEqual(str(template), 'Test Template (v1)')


class TemplateUploadViewTests(APITestCase):
    """Test cases for TemplateUploadView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()
        self.upload_url = reverse('template-upload')

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.regular_user = User.objects.create_user(
            email='regular@example.com',
            password='userpass123'
        )

    def create_docx_file(self, filename='test.docx', with_macros=False):
        """Create a valid DOCX file for testing."""
        from docx import Document

        docx_buffer = io.BytesIO()

        # Create a proper DOCX file using python-docx
        doc = Document()
        doc.add_paragraph('Test content')
        doc.save(docx_buffer)

        if with_macros:
            # Add VBA macro to the DOCX file
            docx_buffer.seek(0)
            with zipfile.ZipFile(docx_buffer, 'a', zipfile.ZIP_DEFLATED) as docx:
                docx.writestr('word/vbaProject.bin', 'VBA_MACRO_CONTENT')

        docx_buffer.seek(0)
        return SimpleUploadedFile(
            filename,
            docx_buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    def test_upload_valid_docx_as_admin(self):
        """Test uploading valid DOCX as admin."""
        self.client.force_authenticate(user=self.admin)

        docx_file = self.create_docx_file()
        data = {
            'file': docx_file,
            'name': 'Content Creator Agreement',
            'category': 'Creator Agreements'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('Template uploaded successfully', response.data['message'])
        self.assertEqual(response.data['template']['name'], 'Content Creator Agreement')
        self.assertEqual(response.data['template']['status'], 'draft')

    def test_template_created_in_database(self):
        """Test that template is created in database."""
        self.client.force_authenticate(user=self.admin)

        docx_file = self.create_docx_file()
        data = {
            'file': docx_file,
            'name': 'Test Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        # Verify template was created
        templates = Template.objects.filter(name='Test Template')
        self.assertTrue(templates.exists())
        template = templates.first()
        self.assertEqual(template.category, 'Test')
        self.assertEqual(template.status, 'draft')

    def test_upload_as_non_admin(self):
        """Test that non-admin users cannot upload templates."""
        self.client.force_authenticate(user=self.regular_user)

        docx_file = self.create_docx_file()
        data = {
            'file': docx_file,
            'name': 'Test Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_upload_file_too_large(self):
        """Test rejecting file larger than 10MB."""
        self.client.force_authenticate(user=self.admin)

        # Create a file larger than 10MB
        large_content = b'x' * (11 * 1024 * 1024)  # 11MB
        large_file = SimpleUploadedFile(
            'large.docx',
            large_content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        data = {
            'file': large_file,
            'name': 'Large Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('10MB limit', str(response.data))

    def test_upload_non_docx_file(self):
        """Test rejecting non-DOCX file."""
        self.client.force_authenticate(user=self.admin)

        txt_file = SimpleUploadedFile(
            'test.txt',
            b'test content',
            content_type='text/plain'
        )

        data = {
            'file': txt_file,
            'name': 'Test Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('Only DOCX format', str(response.data))

    def test_upload_file_with_macros(self):
        """Test rejecting DOCX file with macros."""
        self.client.force_authenticate(user=self.admin)

        docx_file = self.create_docx_file(with_macros=True)
        data = {
            'file': docx_file,
            'name': 'Macro Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('contains macros', response.data['error'])

    def test_checksum_calculated(self):
        """Test that checksum is calculated for uploaded file."""
        self.client.force_authenticate(user=self.admin)

        docx_file = self.create_docx_file()
        data = {
            'file': docx_file,
            'name': 'Test Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('checksum', response.data['template'])
        self.assertEqual(len(response.data['template']['checksum']), 64)  # SHA-256 is 64 chars

    def test_upload_creates_audit_log(self):
        """Test that upload creates audit log."""
        self.client.force_authenticate(user=self.admin)

        docx_file = self.create_docx_file()
        data = {
            'file': docx_file,
            'name': 'Test Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        # Verify audit log
        audit_logs = AuditLog.objects.filter(
            user=self.admin,
            action='template_uploaded'
        )
        self.assertTrue(audit_logs.exists())

    def test_upload_corrupt_docx_file(self):
        """Test uploading corrupt DOCX file."""
        self.client.force_authenticate(user=self.admin)

        # Create a corrupt "DOCX" file (just random bytes)
        corrupt_file = SimpleUploadedFile(
            'corrupt.docx',
            b'This is not a valid DOCX file',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        data = {
            'file': corrupt_file,
            'name': 'Corrupt Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('Invalid DOCX file', response.data['error'])
        self.assertIn('corrupted', response.data['error'].lower())

    def test_macro_detection_creates_audit_log(self):
        """Test that macro detection creates security audit log."""
        self.client.force_authenticate(user=self.admin)

        docx_file = self.create_docx_file(with_macros=True)
        data = {
            'file': docx_file,
            'name': 'Macro Template',
            'category': 'Test'
        }

        response = self.client.post(self.upload_url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)

        # Verify security audit log
        audit_logs = AuditLog.objects.filter(
            user=self.admin,
            action='template_upload_rejected_macros'
        )
        self.assertTrue(audit_logs.exists())


class TemplateListViewTests(APITestCase):
    """Test cases for TemplateListView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()
        self.list_url = reverse('template-list')

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.legal_reviewer = User.objects.create_user(
            email='legal@example.com',
            password='legalpass123'
        )
        self.legal_reviewer.role = 'legal_reviewer'
        self.legal_reviewer.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        self.viewer = User.objects.create_user(
            email='viewer@example.com',
            password='viewerpass123'
        )
        self.viewer.role = 'viewer'
        self.viewer.save()

        # Create test templates with different statuses
        self.draft_template = Template.objects.create(
            name='Draft Template',
            category='Category A',
            file_path='templates/draft.docx',
            checksum_sha256='abc123',
            status='draft',
            uploaded_by=self.admin
        )
        self.approved_template = Template.objects.create(
            name='Approved Template',
            category='Category B',
            file_path='templates/approved.docx',
            checksum_sha256='def456',
            status='approved',
            uploaded_by=self.admin
        )
        self.deprecated_template = Template.objects.create(
            name='Deprecated Template',
            category='Category C',
            file_path='templates/deprecated.docx',
            checksum_sha256='ghi789',
            status='deprecated',
            uploaded_by=self.admin
        )

    def test_admin_sees_all_templates(self):
        """Test that admin can see all templates (all statuses)."""
        self.client.force_authenticate(user=self.admin)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('templates', response.data)
        self.assertEqual(len(response.data['templates']), 3)  # draft, approved, deprecated

    def test_legal_reviewer_sees_all_templates(self):
        """Test that legal reviewer can see all templates (all statuses)."""
        self.client.force_authenticate(user=self.legal_reviewer)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 3)

    def test_content_manager_sees_only_approved_templates(self):
        """Test that content manager can only see approved templates."""
        self.client.force_authenticate(user=self.content_manager)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 1)
        self.assertEqual(response.data['templates'][0]['status'], 'approved')

    def test_viewer_sees_only_approved_templates(self):
        """Test that viewer can only see approved templates."""
        self.client.force_authenticate(user=self.viewer)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 1)
        self.assertEqual(response.data['templates'][0]['status'], 'approved')

    def test_filter_by_status_draft(self):
        """Test filtering templates by draft status."""
        self.client.force_authenticate(user=self.admin)
        response = self.client.get(f'{self.list_url}?status=draft')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 1)
        self.assertEqual(response.data['templates'][0]['status'], 'draft')

    def test_filter_by_status_approved(self):
        """Test filtering templates by approved status."""
        self.client.force_authenticate(user=self.admin)
        response = self.client.get(f'{self.list_url}?status=approved')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 1)
        self.assertEqual(response.data['templates'][0]['status'], 'approved')

    def test_filter_by_category(self):
        """Test filtering templates by category."""
        self.client.force_authenticate(user=self.admin)
        response = self.client.get(f'{self.list_url}?category=Category A')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['templates']), 1)
        self.assertEqual(response.data['templates'][0]['category'], 'Category A')

    def test_template_list_includes_all_fields(self):
        """Test that template list includes all required fields."""
        self.client.force_authenticate(user=self.admin)
        response = self.client.get(self.list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        template_data = response.data['templates'][0]
        self.assertIn('id', template_data)
        self.assertIn('name', template_data)
        self.assertIn('category', template_data)
        self.assertIn('status', template_data)
        self.assertIn('version', template_data)
        self.assertIn('created_at', template_data)


class TemplatePreviewViewTests(APITestCase):
    """Test cases for TemplatePreviewView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.regular_user = User.objects.create_user(
            email='regular@example.com',
            password='userpass123'
        )

        # Create test template with actual DOCX file
        self.template = self.create_template_with_placeholders()

    def create_docx_template_file(self):
        """Create a DOCX template file with Jinja2 placeholders."""
        from docx import Document

        docx_buffer = io.BytesIO()

        # Create a minimal DOCX file with placeholders
        doc = Document()
        doc.add_paragraph('Creator Name: {{ creator_name }}')
        doc.add_paragraph('Amount: {{ amount }}')
        doc.add_paragraph('Date: {{ contract_date }}')
        doc.save(docx_buffer)

        docx_buffer.seek(0)
        return docx_buffer

    def create_template_with_placeholders(self):
        """Create a template with placeholders for testing."""
        # Create DOCX file
        docx_buffer = self.create_docx_template_file()

        # Save to storage
        from django.core.files.base import ContentFile
        file_content = ContentFile(docx_buffer.read())

        template = Template.objects.create(
            name='Test Template',
            category='Test',
            checksum_sha256='abc123',
            uploaded_by=self.admin
        )
        template.file_path.save('test_template.docx', file_content, save=True)

        # Create placeholders
        Placeholder.objects.create(
            template=template,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True,
            position_index=1
        )
        Placeholder.objects.create(
            template=template,
            name='amount',
            display_label='Amount',
            field_type='currency',
            is_required=True,
            position_index=2
        )
        Placeholder.objects.create(
            template=template,
            name='contract_date',
            display_label='Contract Date',
            field_type='date',
            is_required=True,
            position_index=3
        )

        return template

    def test_preview_as_admin_with_dummy_data(self):
        """Test generating preview with dummy data as admin."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.template.id}/preview/'
        response = self.client.post(url, {}, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # Check filename format
        content_disposition = response['Content-Disposition']
        self.assertIn('Test Template_preview_', content_disposition)
        self.assertIn('.docx', content_disposition)

    def test_preview_as_non_admin(self):
        """Test that non-admin users cannot generate previews."""
        self.client.force_authenticate(user=self.regular_user)

        url = f'/api/v1/templates/{self.template.id}/preview/'
        response = self.client.post(url, {}, format='json')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_preview_with_custom_data(self):
        """Test generating preview with custom placeholder data."""
        self.client.force_authenticate(user=self.admin)

        custom_data = {
            'placeholder_data': {
                'creator_name': 'Jane Smith',
                'amount': '75,000',
                'contract_date': '2026-02-01'
            }
        }

        url = f'/api/v1/templates/{self.template.id}/preview/'
        response = self.client.post(url, custom_data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)

    def test_preview_nonexistent_template(self):
        """Test generating preview for nonexistent template."""
        self.client.force_authenticate(user=self.admin)

        url = '/api/v1/templates/99999/preview/'
        response = self.client.post(url, {}, format='json')

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        self.assertIn('Template not found', response.data['error'])

    def test_preview_template_without_placeholders(self):
        """Test generating preview for template without placeholders."""
        self.client.force_authenticate(user=self.admin)

        # Create template without placeholders
        template_no_placeholders = Template.objects.create(
            name='Empty Template',
            category='Test',
            file_path='templates/empty.docx',
            checksum_sha256='xyz789',
            uploaded_by=self.admin
        )

        url = f'/api/v1/templates/{template_no_placeholders.id}/preview/'
        response = self.client.post(url, {}, format='json')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('No placeholders found', response.data['error'])

    def test_preview_creates_audit_log(self):
        """Test that preview creates audit log."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.template.id}/preview/'
        response = self.client.post(url, {}, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify audit log
        audit_logs = AuditLog.objects.filter(
            user=self.admin,
            action='template_preview_generated'
        )
        self.assertTrue(audit_logs.exists())

    def test_dummy_data_generation_for_text_field(self):
        """Test dummy data generation for text field."""
        from agreement_automation.apps.templates.views import generate_dummy_data

        # Create a new template to avoid conflicts
        new_template = Template.objects.create(
            name='Test Template 2',
            category='Test',
            file_path='templates/test2.docx',
            checksum_sha256='xyz789',
            uploaded_by=self.admin
        )

        placeholder = Placeholder.objects.create(
            template=new_template,
            name='test_field',
            display_label='Test Field',
            field_type='text',
            is_required=True
        )

        dummy_data = generate_dummy_data([placeholder])
        self.assertEqual(dummy_data['test_field'], 'Sample Text')

    def test_dummy_data_generation_for_pan_field(self):
        """Test dummy data generation for PAN field."""
        from agreement_automation.apps.templates.views import generate_dummy_data

        # Create a new template to avoid conflicts
        new_template = Template.objects.create(
            name='Test Template 3',
            category='Test',
            file_path='templates/test3.docx',
            checksum_sha256='xyz790',
            uploaded_by=self.admin
        )

        placeholder = Placeholder.objects.create(
            template=new_template,
            name='pan_number',
            display_label='PAN Number',
            field_type='pan_number',
            is_required=True
        )

        dummy_data = generate_dummy_data([placeholder])
        self.assertEqual(dummy_data['pan_number'], 'ABCDE1234F')

    def test_dummy_data_generation_for_gst_field(self):
        """Test dummy data generation for GST field."""
        from agreement_automation.apps.templates.views import generate_dummy_data

        # Create a new template to avoid conflicts
        new_template = Template.objects.create(
            name='Test Template 4',
            category='Test',
            file_path='templates/test4.docx',
            checksum_sha256='xyz791',
            uploaded_by=self.admin
        )

        placeholder = Placeholder.objects.create(
            template=new_template,
            name='gst_number',
            display_label='GST Number',
            field_type='gst_number',
            is_required=True
        )

        dummy_data = generate_dummy_data([placeholder])
        self.assertEqual(dummy_data['gst_number'], '22ABCDE1234F1Z5')

    def test_dummy_data_generation_for_creator_name(self):
        """Test special handling for creator_name field."""
        from agreement_automation.apps.templates.views import generate_dummy_data

        # Use existing placeholder from setUp
        placeholders = self.template.placeholders.filter(name='creator_name')
        dummy_data = generate_dummy_data(placeholders)
        self.assertEqual(dummy_data['creator_name'], 'John Doe')


class TemplateApproveViewTests(APITestCase):
    """Test cases for TemplateApproveView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.legal_reviewer = User.objects.create_user(
            email='legal@example.com',
            password='legalpass123'
        )
        self.legal_reviewer.role = 'legal_reviewer'
        self.legal_reviewer.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        self.viewer = User.objects.create_user(
            email='viewer@example.com',
            password='viewerpass123'
        )
        self.viewer.role = 'viewer'
        self.viewer.save()

        # Create test template with placeholders
        self.template = Template.objects.create(
            name='Test Template',
            category='Test',
            file_path='templates/test.docx',
            checksum_sha256='abc123',
            status='draft',
            uploaded_by=self.admin
        )

        # Add placeholder
        Placeholder.objects.create(
            template=self.template,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True
        )

    def test_admin_can_approve_template(self):
        """Test that admin can approve template."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.template.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('Template approved', response.data['message'])

        # Verify template status changed
        self.template.refresh_from_db()
        self.assertEqual(self.template.status, 'approved')
        self.assertEqual(self.template.approved_by, self.admin)
        self.assertIsNotNone(self.template.approved_at)

    def test_legal_reviewer_can_approve_template(self):
        """Test that legal reviewer can approve template."""
        self.client.force_authenticate(user=self.legal_reviewer)

        url = f'/api/v1/templates/{self.template.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('Template approved', response.data['message'])

        # Verify template status changed
        self.template.refresh_from_db()
        self.assertEqual(self.template.status, 'approved')
        self.assertEqual(self.template.approved_by, self.legal_reviewer)

    def test_content_manager_cannot_approve_template(self):
        """Test that content manager cannot approve template."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/templates/{self.template.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertIn('Access Denied', response.data['error'])

        # Verify template status unchanged
        self.template.refresh_from_db()
        self.assertEqual(self.template.status, 'draft')

    def test_viewer_cannot_approve_template(self):
        """Test that viewer cannot approve template."""
        self.client.force_authenticate(user=self.viewer)

        url = f'/api/v1/templates/{self.template.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertIn('Access Denied', response.data['error'])

    def test_cannot_approve_already_approved_template(self):
        """Test that already approved template cannot be approved again."""
        self.client.force_authenticate(user=self.admin)

        # Approve template first
        self.template.status = 'approved'
        self.template.approved_by = self.admin
        self.template.approved_at = timezone.now()
        self.template.save()

        url = f'/api/v1/templates/{self.template.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('already approved', response.data['error'])

    def test_cannot_approve_deprecated_template(self):
        """Test that deprecated template cannot be approved."""
        self.client.force_authenticate(user=self.admin)

        # Mark template as deprecated
        self.template.status = 'deprecated'
        self.template.save()

        url = f'/api/v1/templates/{self.template.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('deprecated', response.data['error'])

    def test_cannot_approve_template_without_placeholders(self):
        """Test that template without placeholders cannot be approved."""
        self.client.force_authenticate(user=self.admin)

        # Create template without placeholders
        template_no_placeholders = Template.objects.create(
            name='Empty Template',
            category='Test',
            file_path='templates/empty.docx',
            checksum_sha256='xyz789',
            status='draft',
            uploaded_by=self.admin
        )

        url = f'/api/v1/templates/{template_no_placeholders.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('at least one placeholder', response.data['error'])

    def test_approve_nonexistent_template(self):
        """Test approving nonexistent template."""
        self.client.force_authenticate(user=self.admin)

        url = '/api/v1/templates/99999/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        self.assertIn('Template not found', response.data['error'])

    def test_approve_creates_audit_log(self):
        """Test that approval creates audit log."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.template.id}/approve/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify audit log
        audit_logs = AuditLog.objects.filter(
            user=self.admin,
            action='template_approved'
        )
        self.assertTrue(audit_logs.exists())

    def test_template_list_includes_approval_fields(self):
        """Test that template list includes approval fields."""
        self.client.force_authenticate(user=self.admin)

        # Approve template
        self.template.status = 'approved'
        self.template.approved_by = self.admin
        self.template.approved_at = timezone.now()
        self.template.save()

        list_url = reverse('template-list')
        response = self.client.get(list_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        template_data = response.data['templates'][0]
        self.assertIn('approved_by_email', template_data)
        self.assertIn('approved_at', template_data)
        self.assertEqual(template_data['approved_by_email'], self.admin.email)


class TemplateVersioningTests(APITestCase):
    """Test cases for template versioning."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.legal_reviewer = User.objects.create_user(
            email='legal@example.com',
            password='legalpass123'
        )
        self.legal_reviewer.role = 'legal_reviewer'
        self.legal_reviewer.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        # Create approved template v1
        self.template_v1 = Template.objects.create(
            name='Test Template',
            category='Test',
            file_path='templates/test_v1.docx',
            checksum_sha256='abc123',
            version=1,
            status='approved',
            uploaded_by=self.admin,
            approved_by=self.admin,
            approved_at=timezone.now()
        )

        # Add placeholder
        Placeholder.objects.create(
            template=self.template_v1,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True
        )

    def create_docx_file(self, filename='test.docx'):
        """Create a valid DOCX file for testing."""
        from docx import Document

        docx_buffer = io.BytesIO()

        # Create a proper DOCX file using python-docx
        doc = Document()
        doc.add_paragraph('Test content')
        doc.save(docx_buffer)

        docx_buffer.seek(0)
        return SimpleUploadedFile(
            filename,
            docx_buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    def test_create_new_version_from_approved_template(self):
        """Test creating new version from approved template."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.template_v1.id}/new-version/'
        docx_file = self.create_docx_file('test_v2.docx')
        data = {'file': docx_file}

        response = self.client.post(url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('version 2 uploaded', response.data['message'])
        self.assertEqual(response.data['template']['version'], 2)
        self.assertEqual(response.data['template']['status'], 'draft')
        self.assertEqual(response.data['template']['parent_template_id'], self.template_v1.id)

        # Verify new template was created
        new_template = Template.objects.get(id=response.data['template']['id'])
        self.assertEqual(new_template.version, 2)
        self.assertEqual(new_template.parent_template, self.template_v1)
        self.assertEqual(new_template.status, 'draft')

    def test_cannot_create_version_from_draft_template(self):
        """Test that version cannot be created from draft template."""
        self.client.force_authenticate(user=self.admin)

        # Create draft template
        draft_template = Template.objects.create(
            name='Draft Template',
            category='Test',
            file_path='templates/draft.docx',
            checksum_sha256='xyz789',
            version=1,
            status='draft',
            uploaded_by=self.admin
        )

        url = f'/api/v1/templates/{draft_template.id}/new-version/'
        docx_file = self.create_docx_file()
        data = {'file': docx_file}

        response = self.client.post(url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('approved templates', response.data['error'])

    def test_non_admin_cannot_create_new_version(self):
        """Test that non-admin cannot create new version."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/templates/{self.template_v1.id}/new-version/'
        docx_file = self.create_docx_file()
        data = {'file': docx_file}

        response = self.client.post(url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_auto_deprecate_old_version_on_new_approval(self):
        """Test that old version is deprecated when new version is approved."""
        self.client.force_authenticate(user=self.admin)

        # Create version 2
        url = f'/api/v1/templates/{self.template_v1.id}/new-version/'
        docx_file = self.create_docx_file('test_v2.docx')
        data = {'file': docx_file}
        response = self.client.post(url, data, format='multipart')
        template_v2_id = response.data['template']['id']

        # Add placeholder to v2
        template_v2 = Template.objects.get(id=template_v2_id)
        Placeholder.objects.create(
            template=template_v2,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True
        )

        # Verify v1 is still approved
        self.template_v1.refresh_from_db()
        self.assertEqual(self.template_v1.status, 'approved')

        # Approve version 2
        approve_url = f'/api/v1/templates/{template_v2_id}/approve/'
        response = self.client.post(approve_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify v1 is now deprecated
        self.template_v1.refresh_from_db()
        self.assertEqual(self.template_v1.status, 'deprecated')

        # Verify v2 is approved
        template_v2.refresh_from_db()
        self.assertEqual(template_v2.status, 'approved')

    def test_list_all_template_versions(self):
        """Test listing all versions of a template."""
        self.client.force_authenticate(user=self.admin)

        # Create version 2
        url = f'/api/v1/templates/{self.template_v1.id}/new-version/'
        docx_file = self.create_docx_file('test_v2.docx')
        data = {'file': docx_file}
        self.client.post(url, data, format='multipart')

        # List versions
        versions_url = f'/api/v1/templates/{self.template_v1.id}/versions/'
        response = self.client.get(versions_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['versions']), 2)
        self.assertEqual(response.data['template_name'], 'Test Template')

        # Verify versions are in descending order
        self.assertEqual(response.data['versions'][0]['version'], 2)
        self.assertEqual(response.data['versions'][1]['version'], 1)

    def test_legal_reviewer_can_view_versions(self):
        """Test that legal reviewer can view template versions."""
        self.client.force_authenticate(user=self.legal_reviewer)

        versions_url = f'/api/v1/templates/{self.template_v1.id}/versions/'
        response = self.client.get(versions_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)

    def test_content_manager_cannot_view_versions(self):
        """Test that content manager cannot view template versions."""
        self.client.force_authenticate(user=self.content_manager)

        versions_url = f'/api/v1/templates/{self.template_v1.id}/versions/'
        response = self.client.get(versions_url)

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

    def test_version_chain_with_multiple_versions(self):
        """Test version chain with multiple versions."""
        self.client.force_authenticate(user=self.admin)

        # Create version 2
        url = f'/api/v1/templates/{self.template_v1.id}/new-version/'
        docx_file = self.create_docx_file('test_v2.docx')
        response = self.client.post(url, {'file': docx_file}, format='multipart')
        template_v2_id = response.data['template']['id']

        # Approve version 2 (need approved template to create new version)
        template_v2 = Template.objects.get(id=template_v2_id)
        Placeholder.objects.create(
            template=template_v2,
            name='creator_name',
            display_label='Creator Name',
            field_type='text',
            is_required=True
        )
        approve_url = f'/api/v1/templates/{template_v2_id}/approve/'
        self.client.post(approve_url)

        # Create version 3 from version 2
        url = f'/api/v1/templates/{template_v2_id}/new-version/'
        docx_file = self.create_docx_file('test_v3.docx')
        response = self.client.post(url, {'file': docx_file}, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(response.data['template']['version'], 3)
        self.assertEqual(response.data['template']['parent_template_id'], self.template_v1.id)

        # List versions - should show all 3
        versions_url = f'/api/v1/templates/{self.template_v1.id}/versions/'
        response = self.client.get(versions_url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data['versions']), 3)

    def test_new_version_creates_audit_log(self):
        """Test that creating new version creates audit log."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.template_v1.id}/new-version/'
        docx_file = self.create_docx_file()
        data = {'file': docx_file}
        response = self.client.post(url, data, format='multipart')

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)

        # Verify audit log
        audit_logs = AuditLog.objects.filter(
            user=self.admin,
            action='template_version_uploaded'
        )
        self.assertTrue(audit_logs.exists())


class TemplateDeprecateViewTests(APITestCase):
    """Test cases for TemplateDeprecateView."""

    def setUp(self):
        """Set up test client and users."""
        self.client = APIClient()

        self.admin = User.objects.create_user(
            email='admin@example.com',
            password='adminpass123',
            is_staff=True
        )
        self.admin.role = 'admin'
        self.admin.save()

        self.content_manager = User.objects.create_user(
            email='content@example.com',
            password='contentpass123'
        )
        self.content_manager.role = 'content_manager'
        self.content_manager.save()

        # Create approved template
        self.approved_template = Template.objects.create(
            name='Test Template',
            category='Test',
            file_path='templates/test.docx',
            checksum_sha256='abc123',
            version=1,
            status='approved',
            uploaded_by=self.admin,
            approved_by=self.admin,
            approved_at=timezone.now()
        )

        # Create draft template
        self.draft_template = Template.objects.create(
            name='Draft Template',
            category='Test',
            file_path='templates/draft.docx',
            checksum_sha256='xyz789',
            version=1,
            status='draft',
            uploaded_by=self.admin
        )

    def test_admin_can_deprecate_approved_template(self):
        """Test that admin can deprecate approved template."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.approved_template.id}/deprecate/'
        data = {'reason': 'Outdated content'}
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('deprecated', response.data['message'])

        # Verify template status changed
        self.approved_template.refresh_from_db()
        self.assertEqual(self.approved_template.status, 'deprecated')
        self.assertEqual(self.approved_template.deprecated_by, self.admin)
        self.assertIsNotNone(self.approved_template.deprecated_at)
        self.assertEqual(self.approved_template.deprecation_reason, 'Outdated content')

    def test_non_admin_cannot_deprecate_template(self):
        """Test that non-admin cannot deprecate template."""
        self.client.force_authenticate(user=self.content_manager)

        url = f'/api/v1/templates/{self.approved_template.id}/deprecate/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)

        # Verify template status unchanged
        self.approved_template.refresh_from_db()
        self.assertEqual(self.approved_template.status, 'approved')

    def test_cannot_deprecate_already_deprecated_template(self):
        """Test that already deprecated template cannot be deprecated again."""
        self.client.force_authenticate(user=self.admin)

        # Deprecate template first
        self.approved_template.status = 'deprecated'
        self.approved_template.deprecated_by = self.admin
        self.approved_template.deprecated_at = timezone.now()
        self.approved_template.save()

        url = f'/api/v1/templates/{self.approved_template.id}/deprecate/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('already deprecated', response.data['error'])

    def test_cannot_deprecate_draft_template(self):
        """Test that draft template cannot be deprecated."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.draft_template.id}/deprecate/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('draft templates', response.data['error'])

    def test_deprecate_without_reason_uses_default(self):
        """Test deprecating without reason uses default message."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.approved_template.id}/deprecate/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify default reason was used
        self.approved_template.refresh_from_db()
        self.assertEqual(self.approved_template.deprecation_reason, 'Manually deprecated by admin')

    def test_deprecate_nonexistent_template(self):
        """Test deprecating nonexistent template."""
        self.client.force_authenticate(user=self.admin)

        url = '/api/v1/templates/99999/deprecate/'
        response = self.client.post(url)

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        self.assertIn('Template not found', response.data['error'])

    def test_deprecate_creates_audit_log(self):
        """Test that deprecation creates audit log."""
        self.client.force_authenticate(user=self.admin)

        url = f'/api/v1/templates/{self.approved_template.id}/deprecate/'
        data = {'reason': 'Test reason'}
        response = self.client.post(url, data, format='json')

        self.assertEqual(response.status_code, status.HTTP_200_OK)

        # Verify audit log
        audit_logs = AuditLog.objects.filter(
            user=self.admin,
            action='template_deprecated'
        )
        self.assertTrue(audit_logs.exists())
