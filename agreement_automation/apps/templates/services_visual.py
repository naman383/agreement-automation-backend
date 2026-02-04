"""
Services for Visual Template Builder
Handles document processing, region selection, and agreement generation
"""

import hashlib
import mammoth
import re
from docx import Document
from docx.shared import RGBColor, Pt
from io import BytesIO
from typing import Dict, List, Any


class VisualTemplateProcessor:
    """Process uploaded DOCX for visual template builder"""

    @staticmethod
    def process_upload(docx_file) -> Dict[str, Any]:
        """
        Process uploaded DOCX file:
        1. Extract document structure (paragraphs/runs)
        2. Convert to HTML for preview
        3. Calculate checksum
        4. Estimate pages

        Returns dict with processed data
        """
        # Calculate checksum
        docx_file.seek(0)
        checksum = hashlib.sha256(docx_file.read()).hexdigest()
        docx_file.seek(0)

        # Convert to HTML
        result = mammoth.convert_to_html(docx_file)
        html_preview = result.value

        # Extract document structure
        docx_file.seek(0)
        doc = Document(docx_file)
        structure = VisualTemplateProcessor._extract_structure(doc)

        # Estimate pages
        page_count = VisualTemplateProcessor._estimate_pages(doc)

        # Get file size
        docx_file.seek(0, 2)  # Seek to end
        file_size = docx_file.tell()
        docx_file.seek(0)

        return {
            'html_preview': html_preview,
            'document_structure': structure,
            'page_count': page_count,
            'file_size': file_size,
            'checksum': checksum,
            'conversion_messages': [str(m) for m in result.messages]
        }

    @staticmethod
    def _extract_structure(doc: Document) -> Dict[str, Any]:
        """
        Extract paragraph and run structure from DOCX.
        This is used to map selections back to exact positions.
        """
        structure = {
            'paragraph_count': len(doc.paragraphs),
            'paragraphs': []
        }

        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_data = {
                'index': para_idx,
                'text': paragraph.text,
                'runs': [],
                'style': paragraph.style.name if paragraph.style else 'Normal'
            }

            for run_idx, run in enumerate(paragraph.runs):
                run_data = {
                    'index': run_idx,
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'font_size': run.font.size.pt if run.font.size else None,
                }
                para_data['runs'].append(run_data)

            structure['paragraphs'].append(para_data)

        return structure

    @staticmethod
    def _estimate_pages(doc: Document) -> int:
        """Rough page count estimation"""
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        # Average page: ~3000 characters
        return max(1, (total_chars // 3000) + 1)


class RegionSelector:
    """Handle text selection and region creation"""

    @staticmethod
    def create_region_from_selection(selection_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Convert frontend selection to region data.

        selection_data = {
            'selected_text': '___________',
            'html_xpath': '/html/body/p[5]/span[2]',  # Path in HTML
            'visual_coords': {
                'x': 45.5,
                'y': 30.2,
                'width': 20,
                'height': 3,
                'page': 1
            }
        }

        Returns region data ready for PlaceholderRegion model
        """
        # Get visual coordinates with defaults
        visual_coords = selection_data.get('visual_coords', {})

        # Parse HTML path to get paragraph hint
        paragraph_hint = RegionSelector._parse_html_path(selection_data.get('html_xpath', ''))

        return {
            'selected_text': selection_data.get('selected_text', ''),
            'paragraph_index': paragraph_hint,  # Will be refined with structure matching
            'run_index': 0,
            'char_start': 0,
            'char_end': len(selection_data.get('selected_text', '')),
            'page_number': visual_coords.get('page', 1),
            'x_percent': visual_coords.get('x', 0),
            'y_percent': visual_coords.get('y', 0),
            'width_percent': visual_coords.get('width', 0),
            'height_percent': visual_coords.get('height', 0),
        }

    @staticmethod
    def _parse_html_path(xpath: str) -> int:
        """Extract paragraph number from HTML XPath"""
        # Example: /html/body/p[5] → paragraph 4 (0-indexed)
        match = re.search(r'/p\[(\d+)\]', xpath)
        if match:
            return int(match.group(1)) - 1
        return 0

    @staticmethod
    def find_exact_position(doc_structure: Dict, selected_text: str, paragraph_hint: int) -> Dict[str, int]:
        """
        Find exact position of selected text in document structure.
        Uses fuzzy matching around paragraph_hint.
        """
        # Search in hinted paragraph first
        paragraphs = doc_structure['paragraphs']

        if 0 <= paragraph_hint < len(paragraphs):
            para = paragraphs[paragraph_hint]
            position = RegionSelector._find_in_paragraph(para, selected_text)
            if position:
                return {
                    'paragraph_index': paragraph_hint,
                    **position
                }

        # Search nearby paragraphs
        for offset in range(-2, 3):
            idx = paragraph_hint + offset
            if 0 <= idx < len(paragraphs) and idx != paragraph_hint:
                para = paragraphs[idx]
                position = RegionSelector._find_in_paragraph(para, selected_text)
                if position:
                    return {
                        'paragraph_index': idx,
                        **position
                    }

        # Fallback: return hint position
        return {
            'paragraph_index': paragraph_hint,
            'run_index': 0,
            'char_start': 0,
            'char_end': len(selected_text)
        }

    @staticmethod
    def _find_in_paragraph(para: Dict, text: str) -> Dict[str, int] | None:
        """Find text within paragraph runs"""
        for run_idx, run in enumerate(para['runs']):
            run_text = run['text']
            pos = run_text.find(text)
            if pos != -1:
                return {
                    'run_index': run_idx,
                    'char_start': pos,
                    'char_end': pos + len(text)
                }
        return None


class VisualAgreementGenerator:
    """Generate final DOCX with filled data"""

    @staticmethod
    def generate(visual_template, field_values: Dict[str, str]) -> BytesIO:
        """
        Generate agreement DOCX with filled values.

        Args:
            visual_template: VisualTemplate instance
            field_values: {'party_a_name': 'ABC Company', 'date': '2026-02-03'}

        Returns:
            BytesIO with generated DOCX
        """
        # Load original document
        doc = Document(visual_template.original_file.path)

        # Get all placeholders with regions
        placeholders = visual_template.placeholders.all().prefetch_related('regions')

        # Track replacements per paragraph to avoid conflicts
        replacements_by_para = {}

        # Collect all replacements
        for placeholder in placeholders:
            value = field_values.get(placeholder.field_name, '')

            # Format value
            formatted_value = VisualAgreementGenerator._format_value(
                value,
                placeholder.field_type
            )

            # Get all regions for this placeholder
            for region in placeholder.regions.all():
                para_idx = region.paragraph_index
                if para_idx not in replacements_by_para:
                    replacements_by_para[para_idx] = []

                replacements_by_para[para_idx].append({
                    'run_index': region.run_index,
                    'char_start': region.char_start,
                    'char_end': region.char_end,
                    'new_value': formatted_value,
                    'original_text': region.selected_text
                })

        # Apply replacements paragraph by paragraph
        for para_idx, replacements in replacements_by_para.items():
            # Sort by position (reverse) to avoid offset issues
            replacements.sort(key=lambda r: (r['run_index'], r['char_start']), reverse=True)

            try:
                paragraph = doc.paragraphs[para_idx]

                for replacement in replacements:
                    VisualAgreementGenerator._replace_in_run(
                        paragraph,
                        replacement['run_index'],
                        replacement['char_start'],
                        replacement['char_end'],
                        replacement['new_value']
                    )
            except IndexError:
                print(f"Warning: Paragraph {para_idx} not found")
                continue

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output

    @staticmethod
    def _format_value(value: str, field_type: str) -> str:
        """Format value based on field type"""
        if not value:
            return ''

        if field_type == 'currency':
            try:
                num = float(value.replace(',', ''))
                return f"₹{int(num):,}"
            except:
                return value

        elif field_type == 'date':
            try:
                from datetime import datetime
                dt = datetime.fromisoformat(value)
                return dt.strftime('%d-%b-%Y')
            except:
                return value

        elif field_type == 'phone':
            # Format as +91 XXXXX XXXXX
            clean = value.replace(' ', '').replace('+91', '').replace('-', '')
            if len(clean) == 10:
                return f"+91 {clean[:5]} {clean[5:]}"
            return value

        elif field_type == 'pan_number':
            return value.upper()

        elif field_type == 'gst_number':
            return value.upper()

        return str(value)

    @staticmethod
    def _replace_in_run(paragraph, run_index: int, start: int, end: int, new_text: str):
        """Replace text in specific run while preserving formatting"""
        try:
            run = paragraph.runs[run_index]

            # Get original formatting
            font_props = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'size': run.font.size,
                'name': run.font.name,
            }

            # Replace text
            original = run.text
            run.text = original[:start] + new_text + original[end:]

            # Reapply formatting
            run.bold = font_props['bold']
            run.italic = font_props['italic']
            run.underline = font_props['underline']
            if font_props['size']:
                run.font.size = font_props['size']
            if font_props['name']:
                run.font.name = font_props['name']

        except IndexError:
            print(f"Warning: Run {run_index} not found in paragraph")
